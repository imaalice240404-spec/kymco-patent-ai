# -*- coding: utf-8 -*-
import os, json, random, time, datetime, tempfile, io, base64
import streamlit as st
import pandas as pd
import numpy as np
import re
import google.generativeai as genai
from supabase import create_client, Client
import pypdfium2 as pdfium
from docx import Document
import streamlit.components.v1 as components
import plotly.express as px
from PIL import Image, ImageOps

# ==========================================
# ⚙️ 1. 系統初始化與環境設定
# ==========================================
st.set_page_config(page_title="機車專利 AI 戰略分析系統", layout="wide")

def get_config(keys):
    for k in keys:
        try: return st.secrets[k]
        except: continue
    return None

S_URL = get_config(["SUPABASE_URL"])
S_KEY = get_config(["SUPABASE_KEY"])
ADMIN_ID = (get_config(["ADMIN_ID"]) or "admin").upper()

key_pool = []
if get_config(["GOOGLE_API_KEY_1"]): key_pool.append(st.secrets["GOOGLE_API_KEY_1"])
if get_config(["GOOGLE_API_KEY_2"]): key_pool.append(st.secrets["GOOGLE_API_KEY_2"])
if not key_pool and get_config(["GOOGLE_API_KEY"]): key_pool.append(st.secrets["GOOGLE_API_KEY"])

if not all([S_URL, S_KEY, key_pool]):
    st.error("❌ 系統偵測到雲端 Secrets 設定缺失！請檢查 Streamlit 後台設定。")
    st.stop()

SELECTED_G_KEY = random.choice(key_pool)
genai.configure(api_key=SELECTED_G_KEY)
model = genai.GenerativeModel('gemini-2.5-flash', generation_config=genai.types.GenerationConfig(temperature=0.1, top_p=0.8))

@st.cache_resource
def init_supabase() -> Client:
    return create_client(S_URL, S_KEY)
supabase = init_supabase()

# ==========================================
# 🔐 2. 員工職號登入機制
# ==========================================
if 'current_user' not in st.session_state: 
    st.session_state.current_user = None

if not st.session_state.current_user:
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown("<h1 style='text-align: center; color: #1e3a8a;'>🏍️ 專利戰略分析系統</h1>", unsafe_allow_html=True)
        with st.container(border=True):
            st.markdown("### 🔐 內部人員登入")
            job_id_input = st.text_input("請輸入您的員工職號：", placeholder="例如：EMP001")
            if st.button("登入系統", use_container_width=True, type="primary"):
                if job_id_input.strip():
                    st.session_state.current_user = job_id_input.strip().upper()
                    st.rerun()
                else: 
                    st.error("請輸入有效的職號！")
    st.stop()

IS_ADMIN = (st.session_state.current_user == ADMIN_ID)

# ==========================================
# 🧠 3. 高階專利 Prompt 庫
# ==========================================
DETAILED_11_RULES = """
【一、 🚦 FTO 風險判定】
(🔴 紅燈：具威脅 / 🟡 黃燈：需注意 / 🟢 綠燈：已失效。判定原則：若狀態為「公告/核准」或「公開」，絕對不可判定為綠燈！若為「消滅/無效/撤回」才可判為🟢)

【二、📸 技術核心快照】
1. 發明目的：(說明解決傳統弊病) 
2. 核心技術：(說明具體零件結構設計) 
3. 宣稱功效：(說明提升了什麼物理效果)

【三、🏢 研發部門精準派發】
[填入建議部門]。 (分發理由)

【四、🛑 先前技術與妥協分析 (防禦地雷)】
本案欲解決之舊設計缺點：(習用技術缺點)
空間配置限制（破口分析）：(列出獨立項限縮最嚴格之特徵)

【五、🧩 獨立項全要件拆解 (Claim Chart)】
最廣獨立項（請求項1）拆解：(以 1. 2. 3. 逐行條列拆解)
破口：(精準點出最容易被迴避的限制條件)

【六、🪤 附屬項隱藏地雷探測】
(條列出具備具體結構形狀、位置、或參數限制的附屬項)

【七、👁️ 侵權可偵測性評估】
(極易偵測 / 需破壞性拆解，並給出理由)

【八、🕵️‍♂️ 實證功效檢驗 (打假雷達)】
(是否有實體測試數據，或僅為定性描述)

【九、🛡️ 高階迴避設計建議 (防範均等論)】
(提出基於破口的具體修改機構方向)

【十、🧬 技術演進與機構整併雷達】
(分析屬於機構整併或架構重組)

【十一、 🏷️ IPC 分類號分析】
(列出本案的 IPC 分類號，並簡述其代表的技術領域與分類意義)
"""

PROMPT_M1_BATCH = """
你是一位具備 20 年經驗的機車廠資深研發主管兼專利工程師。請嚴格輸出 JSON 格式：
{
  "五大類": "【最高嚴格限制】絕對只能從這 6 個詞彙中挑選：[動力引擎, 車架懸吊, 電裝, 機電, 車體外觀, 其他]。禁止發明新詞。可多選，用半形逗號分隔。",
  "次系統": "自訂 5-8 字的具體系統名",
  "特殊機構": "15字內精準描述其物理改變",
  "達成功效": "20字內描述解決的痛點",
  "核心解法": "用 RD 聽得懂的白話文，精確描述零件之間的連接與作動關係。"
}
"""

PROMPT_M3_SINGLE = f"""
你是一位具備 20 年經驗的機車廠資深研發主管兼專利代理人。請詳細閱讀 PDF。
【🔴 輸出格式要求：純 JSON 格式】
{{
  "rd_card": {{
    "title": "一句話總結", 
    "problem": "傳統缺點", 
    "solution": "本專利特殊結構",
    "risk_check": ["1-1. 獨立項全要件限制A", "1-2. 限制B"],
    "design_avoid_rd": ["針對限制A的迴避方向", "針對限制B的迴避方向"]
  }},
  "vis_data": {{
    "claims": ["1. 獨立項全文...", "2. 依據請求項1..."],
    "components": [ {{"id": "10", "name": "車架"}} ],
    "spec_texts": ["【00xx】段落內容全文"],
    "loophole_quote": "從請求項1中『一字不漏』複製最能代表本案特徵的那一段。⚠️不含習知技術，標點符號需一致。"
  }},
  "ip_report": "請以專業繁體中文撰寫下方【IP報告十一點】內容。"
}}
【重要】：components 元件與標號必須 100% 精確，絕對不可配錯對。
【IP報告結構】：
{DETAILED_11_RULES}
"""

PROMPT_VISION = """
這是一張專利圖。已知元件表：{known_comps}。
請強制找出圖片上「所有肉眼可見的數字標號」！
並精準估算其「數字幾何正中心點」的相對座標 (x_rel, y_rel，範圍 0.000~1.000，請精確到小數點後三位)。
【極度要求】：請仔細掃描，絕對不要漏掉任何一個數字！座標必須對準數字正中心。
嚴格輸出 JSON 格式：{{ "hotspots": [ {{"number": "31", "name": "汽缸頭", "x_rel": 0.452, "y_rel": 0.551}} ] }}
"""

PROMPT_M6_FOREIGN = f"""
這是一份海外專利 PDF。請你以「台灣機車廠資深智權主管」的角色閱讀。
【任務 1】：將其最核心的『請求項 1 (獨立項)』完整翻譯成極度流暢、專業的「台灣繁體中文」。
【任務 2】：用繁體中文撰寫一份 11大天條戰略解析。
【輸出格式：嚴格 JSON】
{{
  "translation": "請在此填寫請求項 1 的繁體中文翻譯...",
  "ip_report": "請以繁體中文撰寫下方【IP報告十一點】並輸出於此。"
}}
【IP報告結構】：
{DETAILED_11_RULES}
"""

# ==========================================
# 🛠️ 4. 共用 CSS/JS (防暈眩 + 狙擊紅圈)
# ==========================================
VIEWER_CSS_JS = """
<style>
    body { margin: 0; font-family: sans-serif; background: #fff; }
    .main-container { display: flex; height: 800px; width: 100%; border: 1px solid #ddd; border-radius: 8px; overflow: hidden; }
    .img-section { flex: 6; position: relative; overflow: auto; background: #f8f9fa; border-right: 2px solid #ddd; padding: 10px; display: flex; justify-content: center; align-items: flex-start;}
    .img-wrapper { position: relative; display: inline-block; }
    .patent-img { max-width: 100%; height: auto; display: block; }
    
    /* 狙擊手準心紅圈 (20px) */
    .hotspot { position: absolute; width: 20px; height: 20px; transform: translate(-50%, -50%); border-radius: 50%; cursor: pointer; transition: 0.2s; border: 2px solid rgba(255,0,0,0.7); background: rgba(255,0,0,0.1); z-index: 10; }
    .hotspot::after { content: ''; position: absolute; top: 50%; left: 50%; transform: translate(-50%, -50%); width: 4px; height: 4px; background: red; border-radius: 50%; }
    .hotspot:hover { background: rgba(255, 0, 0, 0.4); border: 2px solid red; box-shadow: 0 0 10px rgba(255,0,0,0.5); z-index: 50; }
    .hotspot-active { background: rgba(255, 255, 0, 0.6) !important; border: 3px solid red !important; transform: translate(-50%, -50%) scale(1.4); z-index: 50; }
    
    #tooltip { display: none; position: absolute; background: rgba(0, 0, 0, 0.85); color: white; padding: 6px 12px; border-radius: 4px; font-size: 14px; z-index: 100; pointer-events: none; white-space: nowrap; border: 1px solid #fff; }
    .text-section { flex: 4; padding: 20px; overflow-y: auto; font-size: 16px; line-height: 1.8; color: #333; }
    .independent-claim-box { background-color: #fafafa; padding: 15px; border-radius: 8px; border-left: 6px solid #1e3a8a; margin-bottom: 15px; }
    .loophole-highlight { background-color: #ffeb3b; font-weight: bold; color: #b45309; padding: 2px 4px; border-radius: 3px; }
    .dependent-claim { margin-bottom: 15px; color: #555; }
    .comp-text { color: #0284c7; font-weight: bold; cursor: pointer; border-bottom: 1px dashed #0284c7; padding: 0 2px; }
    .highlight-active { background-color: #fef08a; color: #b91c1c; border-radius: 3px; padding: 2px 4px; }
</style>
<script>
    let tooltip;
    document.addEventListener("DOMContentLoaded", () => { tooltip = document.getElementById('tooltip'); });
    
    function hoverImage(num, name) {
        if(!tooltip) tooltip = document.getElementById('tooltip');
        document.onmousemove = e => { tooltip.style.left = (e.pageX + 15) + 'px'; tooltip.style.top = (e.pageY + 15) + 'px'; };
        tooltip.innerHTML = "標號 <b>" + num + "</b> : " + name; 
        tooltip.style.display = 'block'; 
        document.querySelectorAll('.comp-' + num).forEach(el => el.classList.add('highlight-active')); 
    }
    function leaveImage(num) { 
        if(tooltip) tooltip.style.display = 'none'; 
        document.querySelectorAll('.comp-' + num).forEach(el => el.classList.remove('highlight-active')); 
    }
    function hoverText(num) { 
        document.querySelectorAll('.comp-' + num).forEach(el => el.classList.add('highlight-active')); 
        const hs = document.getElementById('hotspot-' + num); if(hs) hs.classList.add('hotspot-active');
    }
    function leaveText(num) { 
        document.querySelectorAll('.comp-' + num).forEach(el => el.classList.remove('highlight-active')); 
        const hs = document.getElementById('hotspot-' + num); if(hs) hs.classList.remove('hotspot-active');
    }
    // 🌟 防暈眩捲動邏輯 (點擊才捲動)
    function clickImageToScrollText(num) { 
        const el = document.querySelector('.comp-' + num);
        if(el) el.scrollIntoView({behavior:'smooth', block:'center'}); 
    }
    function clickTextToScrollImage(num) { 
        const hs = document.getElementById('hotspot-' + num);
        if(hs) hs.scrollIntoView({behavior:'smooth', block:'center'}); 
    }
</script>
"""

# ==========================================
# 🛠️ 5. 後端輔助函數
# ==========================================
def parse_ai_json(text):
    try:
        cln = str(text).replace('```json', '').replace('```', '').strip()
        s, e = cln.find('{'), cln.rfind('}')
        return json.loads(cln[s:e+1]) if s != -1 else {}
    except: return {}

def safe_str(val): return str(val).strip() if pd.notna(val) else ""

def clean_assignee(name):
    name = safe_str(name)
    if not name: return "未知"
    return re.split(r'股份有限公司|有限公司|公司', name)[0].split(' ')[0].strip() if name else "未知"

# 🌟 核心切換：根據當前戰區決定使用哪個資料表
def get_db_table(): 
    return 'patents' if st.session_state.zone_mode == 'TW' else 'global_patents'

DB_COL_MAP = {
    'id': 'ID', 'app_num': '申請號', 'cert_num': '證書號', 'pub_date': '公開公告日', 'app_date': '申請日',
    'assignee': '專利權人', 'title': '專利名稱', 'abstract': '摘要', 'claims': '請求項',
    'legal_status': '案件狀態', 'status': '狀態', 'sys_main': '五大類', 'sys_sub': '次系統',
    'mechanism': '特殊機構', 'effect': '達成功效', 'solution': '核心解法',
    'thumbnail_base64': '代表圖', 'ipc': 'IPC', 'starred_users': '收藏名單', 'user_tags': '用戶標籤',
    'rd_card_json': 'RDJSON', 'vis_data_json': 'VISJSON', 'ip_report_text': 'REPORT'
}

def fetch_patents(status_filter=None):
    try:
        query = supabase.table(get_db_table()).select("*")
        if status_filter: 
            query = query.eq('status', status_filter)
        df = pd.DataFrame(query.execute().data)
        if df.empty: return pd.DataFrame()
        return df.rename(columns=DB_COL_MAP)
    except: return pd.DataFrame()

def crop_white_margins(img):
    try:
        if img.mode != 'RGB': img = img.convert('RGB')
        inv = ImageOps.invert(img)
        bbox = inv.getbbox()
        return img.crop((max(0, bbox[0]-20), max(0, bbox[1]-20), min(img.width, bbox[2]+20), min(img.height, bbox[3]+20))) if bbox else img
    except: return img

def generate_thumbnail_base64(pdf_bytes, page_num=2, max_size=800):
    try:
        doc = pdfium.PdfDocument(pdf_bytes)
        idx = page_num - 1 if 0 <= page_num - 1 < len(doc) else 0 
        pil_image = crop_white_margins(doc[idx].render(scale=2.0).to_pil())
        if hasattr(Image, 'Resampling'): pil_image.thumbnail((max_size, max_size), Image.Resampling.LANCZOS)
        else: pil_image.thumbnail((max_size, max_size), Image.ANTIALIAS)
        buffered = io.BytesIO()
        pil_image.save(buffered, format="JPEG", quality=85)
        return base64.b64encode(buffered.getvalue()).decode()
    except: return None

def create_word_doc(text):
    doc = Document()
    doc.add_heading('專利戰略深度分析報告', 0)
    for para in text.split('\n'):
        if para.strip(): doc.add_paragraph(para.strip())
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def get_ipc4(ipc_str):
    if pd.isna(ipc_str) or not ipc_str or ipc_str == "未知": return []
    res = set()
    for part in re.split(r'[;\|,]', str(ipc_str)):
        match = re.search(r'([A-Z]\d{2}[A-Z])', part.strip().upper())
        if match: res.add(match.group(1))
    return list(res)

def get_patent_type(row):
    cert = str(row.get('證書號', '')).strip().upper()
    app = str(row.get('申請號', '')).strip().upper()
    status = str(row.get('案件狀態', ''))
    if cert.startswith('I') or app.startswith('I') or '公開' in status or '審查' in status: return '發明專利 (I)'
    if cert.startswith('M') or app.startswith('M'): return '新型專利 (M)'
    if cert.startswith('D') or app.startswith('D'): return '設計專利 (D)'
    return '其他'

# ==========================================
# 📊 6. 導覽列與戰略戰區切換引擎
# ==========================================
if 'zone_mode' not in st.session_state: st.session_state.zone_mode = 'TW'
for k in ['rd_card_data','claim_data_t2','scanned_pages','m6_report_data']:
    if k not in st.session_state: st.session_state[k] = {}
if 'target_single_patent' not in st.session_state: st.session_state.target_single_patent = None
if 'ip_report_content' not in st.session_state: st.session_state.ip_report_content = ""
if 'pdf_bytes_main' not in st.session_state: st.session_state.pdf_bytes_main = None
if 'thumbnail_base64' not in st.session_state: st.session_state.thumbnail_base64 = None
if 'target_macro_pool' not in st.session_state: st.session_state.target_macro_pool = pd.DataFrame()
if 'ai_macro_matrix' not in st.session_state: st.session_state.ai_macro_matrix = {}
if 'm5_result' not in st.session_state: st.session_state.m5_result = {}
if 'm6_pdf_bytes' not in st.session_state: st.session_state.m6_pdf_bytes = None

with st.sidebar:
    st.markdown("### 🌍 選擇戰略區域")
    z_choice = st.radio("切換資料庫：", ["🇹🇼 本土防禦 (TW)", "🌎 海外預警 (GLOBAL)"], index=0 if st.session_state.zone_mode == 'TW' else 1)
    new_z = 'TW' if '本土' in z_choice else 'GLOBAL'
    if new_z != st.session_state.zone_mode:
        st.session_state.zone_mode = new_z
        st.rerun()
    
    st.markdown("---")
    st.markdown(f"👤 登入職號：**{st.session_state.current_user}**")
    if st.button("🚪 登出系統", use_container_width=True): 
        st.session_state.clear()
        st.rerun()
    st.markdown("---")

    # 🚨 即時工單通知計數
    open_count = 0
    if IS_ADMIN:
        try: 
            open_count = supabase.table('support_tickets').select('id', count='exact').eq('status', 'OPEN').execute().count or 0
        except: pass
    
    PAGES = ["📥 模組一：探勘匯入", "📊 模組二：研發知識庫", "🕵️ 模組三：單篇深度拆解", "🗺️ 模組四：宏觀地圖", "⚔️ 模組五：組合攻防", "🌍 模組六：海外翻譯中心"]
    if IS_ADMIN: 
        PAGES.append(f"👑 專家工單中心 (🔴 {open_count} 待處理)" if open_count > 0 else "👑 專家工單中心")

    if 'active_tab' not in st.session_state: st.session_state.active_tab = PAGES[1]
    
    # 防止切換頁面時名稱變動導致錯誤
    current_index = 1
    for i, p in enumerate(PAGES):
        if st.session_state.active_tab.startswith(p.split(" ")[0]): current_index = i

    st.session_state.active_tab = st.radio("功能導覽：", PAGES, index=current_index)

    st.markdown("---")
    st.info(f"🗄️ 當前連線：`{get_db_table()}`")
    if st.button("🗑️ 清理當前畫面暫存", use_container_width=True):
        st.session_state.target_single_patent = None
        st.session_state.pdf_bytes_main = None
        for key in ['rd_card_data', 'claim_data_t2', 'scanned_pages', 'm6_report_data']: st.session_state[key] = {}
        st.session_state.ip_report_content = ""
        st.session_state.m6_pdf_bytes = None
        st.rerun()

# ==========================================
# 🚀 模組執行邏輯
# ==========================================
zone_title = "🇹🇼 台灣本土防線" if st.session_state.zone_mode == 'TW' else "🌎 全球海外預警"
st.title(f"🏍️ 機車專利 AI 戰略系統 [{zone_title}]")

# ==========================================
# 👑 管理者工單中心
# ==========================================
if st.session_state.active_tab.startswith("👑 專家"):
    st.header("👑 管理者專屬：專家支援工單中心")
    try:
        tickets = supabase.table('support_tickets').select("*").order('created_at', desc=True).execute().data
        if not tickets: 
            st.success("🎉 目前沒有待處理的工單，大家都很平安！")
        else:
            df_t = pd.DataFrame(tickets)
            open_t = df_t[df_t['status'] == 'OPEN']
            closed_t = df_t[df_t['status'] == 'CLOSED']
            
            t1, t2 = st.tabs([f"🚨 待處理工單 ({len(open_t)})", f"✅ 已結案 ({len(closed_t)})"])
            with t1:
                for _, t in open_t.iterrows():
                    with st.container(border=True):
                        st.markdown(f"### 🎫 工單號: {t['id']} | 專利號: **{t['patent_id']}**")
                        st.caption(f"👤 申請人: {t['job_id']} | 📅 時間: {t['created_at'][:16]}")
                        st.error(f"**🚨 員工疑慮描述：**\n{t['issue_desc']}")
                        
                        ans = st.text_area("✍️ 專業回覆與指導：", key=f"ans_{t['id']}")
                        if st.button("💾 送出回覆並結案", key=f"cls_{t['id']}", type="primary"):
                            supabase.table('support_tickets').update({'admin_reply': ans, 'status': 'CLOSED'}).eq('id', t['id']).execute()
                            st.toast("✅ 工單已結案！")
                            time.sleep(0.5)
                            st.rerun()
            with t2:
                for _, t in closed_t.iterrows():
                    with st.expander(f"✅ [已結案] 專利號: {t['patent_id']} (申請人: {t['job_id']})"):
                        st.write(f"**問題：** {t['issue_desc']}")
                        st.success(f"**回覆：** {t['admin_reply']}")
    except Exception as e: 
        st.error(f"讀取工單失敗: {e}")

# ==========================================
# 📥 模組一：探勘匯入
# ==========================================
elif st.session_state.active_tab == "📥 模組一：探勘匯入":
    st.header(f"1. 資料匯入與狀態更新 (寫入: `{get_db_table()}`)")
    uploaded_excel = st.file_uploader("上傳 TWPAT/Google Patents 匯出的 Excel/CSV", type=["xlsx", "xls", "csv"])

    if uploaded_excel:
        if st.button("🔄 執行資料比對與匯入", type="primary"):
            df = pd.read_csv(uploaded_excel) if uploaded_excel.name.endswith('.csv') else pd.read_excel(uploaded_excel)
            col_map = {
                'title': next((c for c in df.columns if '名稱' in c or '標題' in c or 'title' in c.lower()), None),
                'abs': next((c for c in df.columns if '摘要' in c or 'abstract' in c.lower()), None),
                'claim': next((c for c in df.columns if '範圍' in c or '請求' in c or 'claim' in c.lower()), None),
                'app_num': next((c for c in df.columns if '申請號' in c or 'application' in c.lower()), None),
                'cert_num': next((c for c in df.columns if '證書' in c or '公告' in c or '公開' in c or 'patent' in c.lower() or 'id' in c.lower()), None),
                'app_date': next((c for c in df.columns if '申請日' in c or 'filed' in c.lower()), None),
                'pub_date': next((c for c in df.columns if ('公開日' in c or '公告日' in c or 'pub' in c.lower())), None),
                'assignee': next((c for c in df.columns if '權人' in c or '申請人' in c or 'assignee' in c.lower()), None),
                'status': next((c for c in df.columns if '狀態' in c or 'status' in c.lower()), None),
                'ipc': next((c for c in df.columns if 'IPC' in c.upper()), None)
            }
            
            existing_data = supabase.table(get_db_table()).select("id, app_num, legal_status").execute()
            existing_dict = {d['app_num']: d for d in existing_data.data if d['app_num']}

            new_rows_to_insert = []
            update_count, skip_records = 0, 0
            pb = st.progress(0)
            
            for i, row in df.iterrows():
                app_val = safe_str(row[col_map['app_num']]) if col_map['app_num'] else ""
                cert_val = safe_str(row[col_map['cert_num']]) if col_map['cert_num'] else ""
                new_status = safe_str(row[col_map['status']]) if col_map['status'] else "未知"
                if not app_val and not cert_val: continue 
                
                check_val = app_val if app_val else cert_val

                if check_val in existing_dict:
                    old = existing_dict[check_val]
                    if old['legal_status'] != new_status:
                        upd = {'legal_status': new_status}
                        if "公告" in new_status or "核准" in new_status:
                            upd.update({'cert_num': cert_val, 'rd_card_json': None, 'vis_data_json': None, 'ip_report_text': None})
                        supabase.table(get_db_table()).update(upd).eq('id', old['id']).execute()
                        update_count += 1
                    else: skip_records += 1 
                else:
                    new_rows_to_insert.append({
                        'app_num': app_val, 'cert_num': cert_val,
                        'app_date': safe_str(row[col_map['app_date']]) if col_map['app_date'] else "未知",
                        'pub_date': safe_str(row[col_map['pub_date']]) if col_map['pub_date'] else "未知",
                        'assignee': clean_assignee(safe_str(row[col_map['assignee']])),
                        'title': safe_str(row[col_map['title']]) if col_map['title'] else "無名稱",
                        'abstract': safe_str(row[col_map['abs']]).replace('\n', '')[:500] if col_map['abs'] else "無摘要",
                        'claims': safe_str(row[col_map['claim']]).replace('\n', '')[:500] if col_map['claim'] else "無請求項",
                        'legal_status': new_status, 'status': 'PENDING',
                        'ipc': safe_str(row[col_map['ipc']]) if col_map['ipc'] else "未知",
                        'starred_users': '', 'user_tags': '{}'
                    })
                if i % 10 == 0: pb.progress(min(1.0, (i + 1) / len(df)))
            
            pb.progress(1.0)
            if new_rows_to_insert:
                for i in range(0, len(new_rows_to_insert), 500):
                    supabase.table(get_db_table()).insert(new_rows_to_insert[i:i+500]).execute()
            st.success(f"✅ 同步完成！新增: {len(new_rows_to_insert)} | 更新: {update_count} | 跳過: {skip_records}")

    st.markdown("---")
    st.header("2. AI 批次特徵萃取")
    pend_df = fetch_patents('PENDING')
    if not pend_df.empty:
        bs = st.slider("處理筆數", 1, min(50, len(pend_df)), min(5, len(pend_df)))
        if st.button(f"🤖 啟動探勘管線", type="primary"):
            pb2 = st.progress(0)
            for i, (idx, row) in enumerate(pend_df.head(bs).iterrows()):
                try:
                    res = model.generate_content(f"{PROMPT_M1_BATCH}\n【分析專利】：\n標題:{row['專利名稱']}\n摘要:{row['摘要']}\n請求項:{row['請求項']}").text
                    js = parse_ai_json(res)
                    cats = [c.strip() for c in js.get('五大類', '其他').split(',') if c.strip() in ['動力引擎', '車架懸吊', '電裝', '機電', '車體外觀', '其他']]
                    supabase.table(get_db_table()).update({
                        'sys_main': ', '.join(cats) if cats else '其他', 'sys_sub': js.get('次系統', '未分類'),
                        'mechanism': js.get('特殊機構', ''), 'effect': js.get('達成功效', ''),
                        'solution': js.get('核心解法', ''), 'status': 'COMPLETED'
                    }).eq('id', row['ID']).execute()
                except Exception as e: 
                    supabase.table(get_db_table()).update({'status': 'FAILED'}).eq('id', row['ID']).execute()
                pb2.progress((i + 1) / bs)
                time.sleep(4)
            st.success("✅ 批次解析完成！")
            time.sleep(1)
            st.rerun()

# ==========================================
# 📊 模組二：研發知識庫 (含多用戶隔離收藏)
# ==========================================
elif st.session_state.active_tab == "📊 模組二：研發知識庫":
    df = fetch_patents('COMPLETED')
    if df.empty: 
        st.warning("⚠️ 目前戰區無分析資料，請先至模組一匯入。")
    else:
        df['專利類型'] = df.apply(get_patent_type, axis=1)
        df['IPC4'] = df['IPC'].apply(get_ipc4)
        df['用戶標籤'] = df['用戶標籤'].fillna('{}')
        df['收藏名單'] = df['收藏名單'].fillna('')
        
        cu = st.session_state.current_user
        df['我的標籤'] = df['用戶標籤'].apply(lambda x: json.loads(x or '{}').get(cu, ""))
        df['我已收藏'] = df['收藏名單'].apply(lambda x: cu in str(x).split(','))

        st.header("🔍 研發情報檢索 (R&D Filter Hub)")
        with st.container(border=True):
            col_t1, col_t2 = st.columns([1, 2])
            with col_t1: 
                filter_star = st.checkbox(f"🌟 只顯示 {cu} 的最愛專利")
            with col_t2:
                all_my_tags = sorted(list(set([t.strip() for r in df['我的標籤'] for t in str(r).split(',') if t.strip()])))
                filter_tags = st.multiselect("🏷️ 依「我的專屬標籤」篩選", all_my_tags)

            st.markdown("---")
            col_f1, col_f2, col_f3 = st.columns(3)
            filter_main = col_f1.multiselect("📂 1. 系統分類", pd.Series([c.strip() for r in df['五大類'].astype(str) for c in r.split(',') if c.strip()]).unique())
            filter_sub = col_f2.text_input("⚙️ 2. 次系統 (精確搜尋)")
            search_q = col_f3.text_input("🔑 3. 關鍵字或號碼 (去雜訊搜尋)")

            col_f4, col_f5, col_f6 = st.columns(3)
            filter_company = col_f4.multiselect("🏢 4. 競爭對手", [c for c in df['專利權人'].unique() if c and c != "未知"])
            filter_type = col_f5.multiselect("📜 5. 專利類型", ["發明專利 (I)", "新型專利 (M)", "設計專利 (D)"])
            filter_status = col_f6.multiselect("⚖️ 6. 法律狀態", ["🔴 有效專利 (公告/核准)", "🟡 審查中 (公開)", "🟢 開源/失效 (消滅/撤回/無效)"])

        fdf = df.copy()
        if filter_star: fdf = fdf[fdf['我已收藏'] == True]
        if filter_tags: fdf = fdf[fdf['我的標籤'].apply(lambda x: any(t in str(x) for t in filter_tags))]
        if filter_main: fdf = fdf[fdf['五大類'].astype(str).apply(lambda x: any(m in x for m in filter_main))]
        if filter_sub: fdf = fdf[fdf['次系統'].astype(str).str.contains(filter_sub, na=False)]
        if filter_company: fdf = fdf[fdf['專利權人'].apply(lambda x: any(c in str(x) for c in filter_company))]
        if filter_type: fdf = fdf[fdf['專利類型'].isin(filter_type)]
        
        if filter_status:
            stat_mask = pd.Series([False]*len(fdf), index=fdf.index)
            for s in filter_status:
                if "有效" in s: stat_mask |= fdf['案件狀態'].astype(str).str.contains("公告|核准")
                elif "審查中" in s: stat_mask |= fdf['案件狀態'].astype(str).str.contains("公開") & ~fdf['案件狀態'].astype(str).str.contains("公告|核准")
                elif "開源" in s: stat_mask |= fdf['案件狀態'].astype(str).str.contains("消滅|撤回|無效|核駁")
            fdf = fdf[stat_mask]

        if search_q:
            qc = re.sub(r'[^a-zA-Z0-9]', '', search_q).upper()
            fdf = fdf[fdf.astype(str).apply(lambda x: x.str.contains(search_q, case=False)).any(axis=1) | 
                      fdf['證書號'].astype(str).str.replace(r'[^a-zA-Z0-9]', '', regex=True).str.upper().str.contains(qc) |
                      fdf['申請號'].astype(str).str.replace(r'[^a-zA-Z0-9]', '', regex=True).str.upper().str.contains(qc)]

        st.info(f"✨ 檢索出 **{len(fdf)}** 筆專利。")

        col_act1, col_act2 = st.columns(2)
        with col_act1:
            if st.button("🗺️ 將下方專利送往【宏觀地圖】分析", use_container_width=True, type="primary"):
                st.session_state.target_macro_pool = fdf
                st.session_state.ai_macro_matrix = None
                st.session_state.active_tab = PAGES[3]
                st.rerun()
        with col_act2:
            if st.button("⚔️ 切換至【組合攻防】分析", use_container_width=True, type="secondary"):
                st.session_state.active_tab = PAGES[4]
                st.rerun()
        st.markdown("---")

        for _, p in fdf.iterrows():
            did = p['證書號'] if p['證書號'] else p['申請號']
            with st.container(border=True):
                cc, ct, ccont, cm = st.columns([0.5, 2.5, 5.5, 2])
                with cc: 
                    st.write("")
                    st.checkbox("選取", key=f"chk_{did}", label_visibility="collapsed")
                with ct:
                    st.write("")
                    if p.get('代表圖') and len(str(p.get('代表圖'))) > 100: 
                        st.image(f"data:image/jpeg;base64,{p['代表圖']}", use_container_width=True)
                    else: 
                        st.markdown("<div style='border:1px dashed #ccc; height:180px; display:flex; align-items:center; justify-content:center; color:#999; border-radius:8px;'>🖼️ 無代表圖</div>", unsafe_allow_html=True)
                
                with ccont:
                    s_disp = "🌟" if p['我已收藏'] else ""
                    st.markdown(f"#### {s_disp} [{did}] {p['專利名稱']}")
                    st.caption(f"🏢 {p['專利權人']} ｜ 📅 日期: {p['公開公告日']} ｜ 🏷️ {p['專利類型']}")
                    c1, c2, c3 = st.columns([2, 1.5, 2])
                    c1.info(f"📂 **分類**：\n{p['五大類']} ➡️ {p['次系統']}")
                    c2.warning(f"⚙️ **機構**：\n{p['特殊機構']}")
                    c3.error(f"🎯 **功效**：\n{p['達成功效']}")
                    if p.get('核心解法'): 
                        st.markdown(f"<div style='border-left: 4px solid #ddd; padding-left: 15px; color: #555; line-height: 1.6; font-size: 15px;'>💡 **解法：**{p['核心解法']}</div>", unsafe_allow_html=True)

                with cm:
                    st.write("")
                    if st.button("取消我的收藏" if p['我已收藏'] else "⭐ 加入我的最愛", key=f"star_{did}", use_container_width=True):
                        arr = [u for u in str(p['收藏名單']).split(',') if u]
                        if p['我已收藏'] and cu in arr: arr.remove(cu)
                        elif not p['我已收藏'] and cu not in arr: arr.append(cu)
                        supabase.table(get_db_table()).update({'starred_users': ",".join(arr)}).eq('id', p['ID']).execute()
                        st.rerun()

                    nt = st.text_input("專屬標籤", value=p['我的標籤'], placeholder="#專案A", key=f"tags_{did}")
                    if nt != p['我的標籤']:
                        if st.button("💾 儲存標籤", key=f"savetag_{did}", use_container_width=True):
                            td = json.loads(p['用戶標籤']) if p['用戶標籤'] else {}
                            td[cu] = nt
                            supabase.table(get_db_table()).update({'user_tags': json.dumps(td, ensure_ascii=False)}).eq('id', p['ID']).execute()
                            st.rerun()

                    st.markdown("<br>", unsafe_allow_html=True)
                    if st.button("📄 進入深度拆解", key=f"btn_s_{did}", use_container_width=True, type="primary"):
                        st.session_state.target_single_patent = p.to_dict()
                        st.session_state.pdf_bytes_main = None 
                        for key in ['rd_card_data', 'claim_data_t2', 'scanned_pages']: st.session_state[key] = {}
                        st.session_state.ip_report_content, st.session_state.thumbnail_base64 = "", None
                        st.session_state.active_tab = PAGES[2]
                        st.rerun()

                if IS_ADMIN:
                    with st.expander("✏️ 手動修改狀態"):
                        ce1, ce2, ce3 = st.columns([2, 2, 1])
                        with ce1: 
                            nl = st.text_input("法律狀態", value=p['案件狀態'], key=f"leg_{did}")
                        with ce2: 
                            so = ["COMPLETED", "PENDING", "FAILED"]
                            ns = st.selectbox("分析狀態", so, index=so.index(p['狀態']) if p['狀態'] in so else 0, key=f"sys_{did}")
                        with ce3:
                            st.write("")
                            if st.button("💾", key=f"save_{did}"):
                                supabase.table(get_db_table()).update({'legal_status': nl, 'status': ns}).eq('id', p['ID']).execute()
                                st.rerun()

# ==========================================
# 🕵️ 模組三：單篇深度拆解 (包含 11大天條與高精度視覺)
# ==========================================
elif st.session_state.active_tab == "🕵️ 模組三：單篇深度拆解":
    t = st.session_state.target_single_patent
    if not t: 
        st.warning("👈 請先從模組二選擇一篇專利。")
    else:
        db_id, did = t.get('ID'), (t.get('證書號') or t.get('申請號'))
        st.header(f"🕵️ 深度拆解：[{did}] {t.get('專利名稱')}")
        st.markdown(f"**🏢 權利人：** {t.get('專利權人')} | **📅 公開日：** {t.get('公開公告日')} ｜ 🏷️ 類型：{t.get('專利類型')}")
        st.markdown("---")

        if not st.session_state.rd_card_data:
            res = supabase.table(get_db_table()).select("rd_card_json, vis_data_json, ip_report_text").eq('id', db_id).execute().data
            if res and res[0].get('rd_card_json'):
                st.toast("✅ 自動載入歷史報告！")
                st.session_state.rd_card_data = res[0].get('rd_card_json')
                st.session_state.claim_data_t2 = res[0].get('vis_data_json')
                st.session_state.ip_report_content = res[0].get('ip_report_text')

        if not st.session_state.rd_card_data:
            st.info("💡 此專利尚未深度拆解，請上傳 PDF 說明書。")
            with st.container(border=True):
                up_pdf = st.file_uploader("📂 上傳 PDF", type=["pdf"])
                pg_num = st.number_input("🖼️ 代表圖頁碼", min_value=1, value=2)
                if up_pdf and st.button("🚀 啟動 11 大天條超級解析", type="primary", use_container_width=True):
                    st.session_state.pdf_bytes_main = up_pdf.getvalue()
                    tb64 = generate_thumbnail_base64(st.session_state.pdf_bytes_main, page_num=pg_num)
                    with st.spinner("🧠 正在搜索全文並套用 11 大天條 (約 40 秒)..."):
                        try:
                            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp: 
                                tmp.write(st.session_state.pdf_bytes_main)
                                tp = tmp.name
                            gf = genai.upload_file(tp)
                            
                            res = model.generate_content([gf, PROMPT_M3_SINGLE])
                            js = parse_ai_json(res.text) 
                            
                            st.session_state.rd_card_data = js.get("rd_card", {})
                            st.session_state.claim_data_t2 = js.get("vis_data", {})
                            st.session_state.ip_report_content = js.get("ip_report", "")
                            
                            supabase.table(get_db_table()).update({
                                'rd_card_json': st.session_state.rd_card_data,
                                'vis_data_json': st.session_state.claim_data_t2,
                                'ip_report_text': st.session_state.ip_report_content,
                                'thumbnail_base64': tb64
                            }).eq('id', db_id).execute()
                            
                            try: genai.delete_file(gf.name)
                            except: pass 
                            os.remove(tp)
                            st.rerun()
                        except Exception as e: 
                            st.error(f"分析發生異常：{e}")

        if st.session_state.rd_card_data:
            t_rd, t_ip = st.tabs(["🧑‍💻 Tab 1: 研發迴避大屏", "⚖️ Tab 2: 智權 11 大天條中心"])
            
            with t_rd:
                rd = st.session_state.rd_card_data
                c1, c2, c3 = st.columns([1.5, 2, 1.5])
                with c1:
                    with st.container(border=True, height=480):
                        st.markdown(f"#### 🎯 研發戰略看板\n**{rd.get('title', '')}**")
                        st.markdown(f"**🔥 解決痛點：**\n\n{rd.get('problem', '')}\n\n**💡 核心解法：**\n\n{rd.get('solution', '')}")
                with c2:
                    with st.container(border=True, height=480):
                        st.markdown("#### 🛡️ 獨立項全要件檢核")
                        st.caption("全要件原則：符合下方所有特徵，則侵權風險極高。")
                        ck_cnt = 0
                        r_list = rd.get('risk_check', [])
                        for i, r in enumerate(r_list):
                            if st.checkbox(str(r), key=f"rc_{i}"): ck_cnt += 1
                        st.markdown("<br>", unsafe_allow_html=True)
                        if r_list:
                            if ck_cnt == len(r_list): 
                                st.markdown("<div style='padding:10px; background-color:#ffebee; color:#c62828; border-radius:5px;'><b>⚠️ 高度侵權風險！</b></div>", unsafe_allow_html=True)
                            else: 
                                st.markdown("<div style='padding:10px; background-color:#e8f5e9; color:#2e7d32; border-radius:5px;'><b>🎉 文義迴避成功。</b></div>", unsafe_allow_html=True)
                with c3:
                    with st.container(border=True, height=480):
                        st.markdown("#### 🛡️ 高階迴避建議")
                        for a in rd.get('design_avoid_rd', []): 
                            st.markdown(f"✅ {a}")

                st.markdown("---")
                st.markdown("### 🎯 終極雙向連動大屏")
                if not st.session_state.pdf_bytes_main:
                    sup = st.file_uploader("📂 補傳 PDF 解鎖互動圖面", type=["pdf"])
                    if sup: 
                        st.session_state.pdf_bytes_main = sup.getvalue()
                        st.rerun()
                else:
                    doc = pdfium.PdfDocument(st.session_state.pdf_bytes_main)
                    c_pg, c_rt, c_bt = st.columns([1, 1.5, 1.5])
                    pg = c_pg.number_input("📄 頁碼", 1, len(doc), 2)
                    rot = c_rt.radio("🔄 旋轉", [0, 90, 180, 270], horizontal=True)
                    
                    pil = doc[pg-1].render(scale=2.0).to_pil()
                    if rot != 0: pil = pil.rotate(-rot, expand=True, fillcolor='white')
                    img_b = io.BytesIO(); crop_white_margins(pil).save(img_b, 'JPEG')
                    uri = f"data:image/jpeg;base64,{base64.b64encode(img_b.getvalue()).decode()}"
                    sk = f"{pg}_{rot}"
                    cd = st.session_state.claim_data_t2
                    
                    with c_bt:
                        st.write("")
                        if sk not in st.session_state.scanned_pages:
                            if st.button("🔍 鎖定圖片標號座標", use_container_width=True):
                                with st.spinner("AI 像素精準估算中..."):
                                    try:
                                        pv = PROMPT_VISION.format(known_comps=json.dumps(cd.get("components", []), ensure_ascii=False))
                                        res_vis = model.generate_content([crop_white_margins(pil), pv]).text
                                        st.session_state.scanned_pages[sk] = parse_ai_json(res_vis).get("hotspots", [])
                                        st.rerun()
                                    except Exception as e: 
                                        st.error("視覺失敗")
                        else: 
                            st.success("⚡ 座標已鎖定！體驗點擊雙向連動。")

                    hs_html = ""
                    for s in st.session_state.scanned_pages.get(sk, []):
                        sn, snm = str(s.get('number','')), str(s.get('name',''))
                        hs_html += f"<div class='hotspot' id='hotspot-{sn}' style='left:{s.get('x_rel',0)*100}%; top:{s.get('y_rel',0)*100}%;' onmouseover=\"hoverImage('{sn}', '{snm}')\" onmouseout=\"leaveImage('{sn}')\" onclick=\"clickImageToScrollText('{sn}')\"></div>"
                    
                    ct_html = ""
                    lq = str(cd.get("loophole_quote", ""))
                    for i, line in enumerate(cd.get("claims", [])):
                        L = str(line)
                        if i == 0 and lq and lq in L: 
                            L = L.replace(lq, f'<mark class="loophole-highlight">{lq}</mark>')
                        for c in cd.get("components", []):
                            if isinstance(c, dict) and c.get("id") and c.get("name"):
                                cid, cnm = str(c.get("id")), str(c.get("name"))
                                rep = f'<span class="comp-text comp-{cid}" onmouseover="hoverText(\'{cid}\')" onmouseout="leaveText(\'{cid}\')" onclick="clickTextToScrollImage(\'{cid}\')">{cnm} ({cid})</span>'
                                L = L.replace(f"{cnm} ({cid})", rep).replace(cnm, rep)
                        
                        if i == 0: 
                            ct_html += f"<div class='independent-claim-box'>{L}</div>"
                        else: 
                            ct_html += f"<p class='dependent-claim'>{L}</p>"
                    
                    html_skeleton = f"""
                    <!DOCTYPE html>
                    <html>
                    <head>{VIEWER_CSS_JS}</head>
                    <body>
                        <div class='main-container'>
                            <div class='img-section'>
                                <div class='img-wrapper'>
                                    <img src='{uri}' class='patent-img'>{hs_html}
                                </div>
                                <div id='tooltip'></div>
                            </div>
                            <div class='text-section'>
                                <div style='font-size:18px; font-weight:bold; color:#1e3a8a; margin-bottom:15px; position:sticky; top:0; background:white; z-index:10;'>📜 請求項對應 (點擊文字捲動至圖面)</div>
                                {ct_html}
                            </div>
                        </div>
                    </body>
                    </html>
                    """
                    components.html(html_skeleton, height=820, scrolling=False)

            with t_ip:
                st.markdown("### ⚖️ 智權法務深度報告")
                r1, r2 = st.tabs(["📄 智權戰略深度報告 (11大天條)", "⚖️ 請求項文義比對 (三視窗)"])
                with r1:
                    c_dl, c_dr = st.columns([3, 1])
                    c_dl.markdown("以下為嚴格遵守「智權審查 11 大天條」生成的實務報告：")
                    c_dr.download_button("📥 下載 Word 報告", data=create_word_doc(st.session_state.ip_report_content), file_name=f"IP_Report_{did}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                    with st.container(height=650, border=True): 
                        st.markdown(st.session_state.ip_report_content)
                with r2:
                    if not st.session_state.pdf_bytes_main: 
                        st.info("請先補傳 PDF。")
                    else:
                        cl_list = st.session_state.claim_data_t2.get("components", [])
                        if cl_list:
                            opts = {f"[{c.get('id','')}] {c.get('name','')}": c for c in cl_list if isinstance(c, dict) and c.get('id')}
                            if opts:
                                ac = opts[st.selectbox("🎯 目標元件：", list(opts.keys()))]
                                cw1, cw2 = st.columns(2)
                                with cw1:
                                    st.markdown("### 🧩 獨立項文義")
                                    with st.container(height=350, border=True):
                                        for L in st.session_state.claim_data_t2.get("claims", []):
                                            Ls = str(L)
                                            if ac['name'] in Ls: 
                                                # 安全替換，避免 SyntaxError
                                                highlighted_str = Ls.replace(ac['name'], f"<span style='background-color:#fff3cd; font-weight:bold; color:#856404; padding:2px 4px; border-radius:3px;'>{ac['name']}</span>")
                                                st.markdown(f"<div style='padding: 8px; border-bottom: 1px dashed #eee;'>{highlighted_str}</div>", unsafe_allow_html=True)
                                            else: 
                                                st.markdown(f"<div style='padding: 8px; border-bottom: 1px dashed #eee; color: #555;'>{Ls}</div>", unsafe_allow_html=True)
                                    st.markdown("### 🖼️ 專利圖面")
                                    pdf_ip = pdfium.PdfDocument(st.session_state.pdf_bytes_main)
                                    pg_ip = st.number_input("對照頁碼", 1, len(pdf_ip), 1)
                                    with st.container(height=450, border=True): 
                                        st.image(pdf_ip[pg_ip-1].render(scale=2.0).to_pil(), use_container_width=True)
                                
                                with cw2:
                                    st.markdown("### 📖 說明書具體限制")
                                    with st.container(height=895, border=True):
                                        fts = [t for t in st.session_state.claim_data_t2.get('spec_texts', []) if ac['name'] in str(t)]
                                        if not fts: 
                                            st.warning("未找到說明。")
                                        else:
                                            for t in fts: 
                                                # 安全替換，避免 SyntaxError
                                                highlighted_text = str(t).replace(ac['name'], f"<mark style='background-color:#cce5ff; color:#004085; font-weight:bold; padding:2px; border-radius:3px;'>{ac['name']}</mark>")
                                                st.markdown(f"<div style='background: #f8f9fa; padding: 10px; border-left: 4px solid #007bff; margin-bottom: 10px;'>{highlighted_text}</div>", unsafe_allow_html=True)

            with st.expander("🚨 AI 解析結果不滿意？呼叫管理者支援"):
                issue = st.text_area("描述您需要協助的部分：")
                if st.button("📨 送出支援工單", type="primary"):
                    if issue.strip():
                        supabase.table("support_tickets").insert({"patent_id": did, "job_id": st.session_state.current_user, "issue_desc": issue.strip()}).execute()
                        st.success("✅ 工單已送出！")
                    else: 
                        st.error("請填寫描述。")

# --- 模組四：宏觀分析 ---
elif st.session_state.active_tab == "🗺️ 模組四：宏觀地圖":
    st.header("🗺️ 傳統專利大數據分析")
    if st.session_state.target_macro_pool.empty: 
        st.warning("請先從模組二過濾專利並點擊【送往宏觀地圖】按鈕。")
    else:
        dfm = st.session_state.target_macro_pool
        st.success(f"✅ 已載入 {len(dfm)} 筆專利進行宏觀分析！")
        
        t1, t2, t3, t4, t5 = st.tabs(["🏢 競爭者佈局", "📈 演進趨勢", "🎯 IPC 熱區分析", "🧠 AI 功效矩陣", "👑 核心地雷探勘"])
        
        with t1:
            st.markdown("### 🏢 專利權人火力佈局")
            top_app = dfm['專利權人'].value_counts().reset_index().head(10)
            top_app.columns = ['公司名稱', '專利數量']
            fig_bar = px.bar(top_app, x='專利數量', y='公司名稱', orientation='h', color='專利數量', color_continuous_scale='Blues')
            fig_bar.update_layout(yaxis={'categoryorder':'total ascending'})
            st.plotly_chart(fig_bar, use_container_width=True)

        with t2:
            st.markdown("### 📈 歷年趨勢")
            dfm['年份'] = dfm['公開公告日'].astype(str).str[:4]
            yt = dfm['年份'].value_counts().reset_index().sort_values('年份')
            yt.columns = ['年份', '專利數量']
            fig_line = px.line(yt[yt['年份'].str.isnumeric()], x='年份', y='專利數量', markers=True)
            st.plotly_chart(fig_line, use_container_width=True)

        with t3:
            st.markdown("### 🎯 IPC 熱區")
            if 'IPC4' in dfm.columns:
                all_ipc_flat = [ipc for sublist in dfm['IPC4'] for ipc in sublist]
                if all_ipc_flat:
                    ipc_d = pd.Series(all_ipc_flat).value_counts().reset_index().head(15)
                    ipc_d.columns = ['IPC四階', '數量']
                    fig_pie = px.pie(ipc_d, values='數量', names='IPC四階', hole=0.4)
                    st.plotly_chart(fig_pie, use_container_width=True)

        with t4:
            st.markdown("### 🧠 AI 自動生成：技術功效矩陣")
            ac = st.slider("分析筆數", 1, min(len(dfm), 30), min(len(dfm), 15))
            if st.button("🚀 啟動矩陣引擎", type="primary"):
                with st.spinner("掃描戰略地圖..."):
                    try:
                        sdf = dfm.head(ac)
                        p_data = "".join([f"[{row['證書號'] or row['申請號']}] {row['專利名稱']} | 機構：{row['特殊機構']} | 功效：{row['達成功效']}\n" for _, row in sdf.iterrows()])
                        
                        prompt_matrix = "\n".join([
                            "分析機車專利資料，輸出 JSON 格式。",
                            "矩陣維度X (達成功效): [\"提升散熱與冷卻\", \"提升動力效率\", \"結構輕量化\", \"降低震動\", \"改善耐用度\", \"降低成本\"]。",
                            "矩陣維度Y (技術手段): [\"汽缸與散熱片\", \"活塞曲軸\", \"氣門系統\", \"機油道水套\", \"燃油點火\", \"引擎外殼\", \"懸吊\", \"電控\"]。",
                            "{",
                            "  \"matrix\": [{\"專利號\": \"XXX\", \"技術手段\": \"選項\", \"達成功效\": \"選項\"}],",
                            "  \"top_patents\": [{\"專利號\": \"XXX\", \"專利名稱\": \"XXX\", \"威脅度\": \"🔴極高/🟡中等\", \"入選理由\": \"...\"}]",
                            "}",
                            f"資料：\n{p_data}"
                        ])
                        res = model.generate_content(prompt_matrix)
                        st.session_state.ai_macro_matrix = parse_ai_json(res.text)
                        st.success("解析完成！")
                    except Exception as e: st.error(f"失敗：{e}")

            if st.session_state.ai_macro_matrix:
                mat = st.session_state.ai_macro_matrix.get("matrix", [])
                if mat:
                    fig_heat = px.density_heatmap(pd.DataFrame(mat), y='技術手段', x='達成功效', text_auto=True, color_continuous_scale='Reds')
                    st.plotly_chart(fig_heat, use_container_width=True)

        with t5:
            st.markdown("### 👑 核心地雷探勘 (Killer Patents)")
            if st.session_state.ai_macro_matrix:
                for p in st.session_state.ai_macro_matrix.get("top_patents", []):
                    with st.container(border=True):
                        c = "red" if "高" in p.get("威脅度", "") else "orange"
                        st.markdown(f"#### 🎯 [{p.get('專利號')}] {p.get('專利名稱')}")
                        st.markdown(f"**威脅度：** <span style='color:{c};font-weight:bold;'>{p.get('威脅度')}</span><br>**洞察：** {p.get('入選理由')}", unsafe_allow_html=True)

# --- 模組五：組合核駁 ---
elif st.session_state.active_tab == "⚔️ 模組五：組合攻防":
    st.header("⚔️ 組合核駁與 TSM 攻防")
    df_m5 = fetch_patents('COMPLETED')
    if df_m5.empty or len(df_m5) < 2:
        st.warning("請先由模組一匯入充足引證案。")
    else:
        opts = [f"[{r['證書號'] or r['申請號']}] {r['專利名稱']}" for _, r in df_m5.iterrows()]
        opt_map = {f"[{r['證書號'] or r['申請號']}] {r['專利名稱']}": r for _, r in df_m5.iterrows()}

        c1, c2 = st.columns(2)
        r_a = c1.selectbox("📄 引證一", opts, index=0)
        r_b = c2.selectbox("📄 引證二", opts, index=min(1, len(opts)-1))

        with st.container(border=True):
            tf = st.text_area("1. 🛡️ 本案爭點特徵", height=100)
            cc1, cc2 = st.columns(2)
            ra_d = cc1.text_area("2. 📄 引證一具體揭露", height=150)
            rb_d = cc2.text_area("3. 📄 引證二具體揭露", height=150)
            el = st.text_area("4. ⚔️ 審查委員結合邏輯", height=100)

        if st.button("🚀 啟動 TSM 分析", type="primary"):
            if not tf or not ra_d or not rb_d: st.error("請填寫完整")
            else:
                with st.spinner("沙盤推演中..."):
                    ref_a, ref_b = opt_map[r_a], opt_map[r_b]
                    pm5 = "\n".join([
                        "【角色設定】：資深機車專利代理人進行 TSM 攻防。",
                        f"- 爭點：{tf}\n- 引證A揭露：{ra_d}\n- 引證B揭露：{rb_d}\n- 審委邏輯：{el}",
                        f"[A背景] 摘要：{ref_a['摘要']} | 核心解法：{ref_a['核心解法']}",
                        f"[B背景] 摘要：{ref_b['摘要']} | 核心解法：{ref_b['核心解法']}",
                        "【嚴格輸出 JSON】：{\"delta_feature\":\"...\",\"attack_argument\":{\"conclusion\":\"...\",\"field_problem_match\":\"...\",\"motivation_to_combine\":\"...\"},\"defense_argument\":{\"conclusion\":\"...\",\"teaching_away\":\"...\",\"hindsight_warning\":\"...\"}}"
                    ])
                    try:
                        res = model.generate_content(pm5)
                        st.session_state.m5_result = parse_ai_json(res.text)
                        st.success("✅ TSM 雙向攻防推演完成！")
                    except Exception as e: st.error("失敗")

        if st.session_state.m5_result:
            m5 = st.session_state.m5_result
            st.markdown(f"### 🎯 差異特徵\n<div style='background-color:#e0f2fe; padding:15px; border-radius:8px;'><b>{m5.get('delta_feature', '')}</b></div>", unsafe_allow_html=True)
            t_def, t_atk = st.tabs(["🛡️ 防禦視角", "⚔️ 攻擊視角"])
            with t_def:
                st.error(f"**阻礙結合：** {m5.get('defense_argument', {}).get('teaching_away', '')}")
                st.warning(f"**後見之明：** {m5.get('defense_argument', {}).get('hindsight_warning', '')}")
            with t_atk:
                st.info(f"**共通性：** {m5.get('attack_argument', {}).get('field_problem_match', '')}")
                st.success(f"**結合動機：** {m5.get('attack_argument', {}).get('motivation_to_combine', '')}")

# --- 🌍 模組六：海外翻譯中心 ---
elif st.session_state.active_tab == "🌍 模組六：海外翻譯中心":
    st.header("🌍 海外專利雙視窗翻譯對照中心")
    st.info("此模組為海外專利臨時解析區。")
    
    up6 = st.file_uploader("📂 載入日/美/歐專利 PDF", type=["pdf"])
    if up6 and st.button("🚀 啟動 AI 繁中智能編譯", type="primary"):
        with st.spinner("跨國智權解析與翻譯中 (約 40 秒)..."):
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp: 
                    tmp.write(up6.getvalue())
                    tp = tmp.name
                gf = genai.upload_file(tp)
                
                res = model.generate_content([gf, PROMPT_M6_FOREIGN]).text
                st.session_state.m6_report_data = parse_ai_json(res)
                st.session_state.m6_pdf_bytes = up6.getvalue()
                
                try: genai.delete_file(gf.name)
                except: pass
                os.remove(tp)
                st.rerun()
            except Exception as e: 
                st.error(f"編譯失敗: {e}")

    if st.session_state.m6_pdf_bytes and st.session_state.m6_report_data:
        st.markdown("---")
        c_l, c_r = st.columns([1.2, 1])
        with c_l:
            st.subheader("📄 原始專利 PDF")
            m6_doc = pdfium.PdfDocument(st.session_state.m6_pdf_bytes)
            pg = st.number_input("原件頁碼", 1, len(m6_doc), 1)
            st.image(m6_doc[pg-1].render(scale=2.0).to_pil(), use_container_width=True)
        with c_r:
            st.subheader("🧠 AI 繁體中文對照區")
            ta, tb = st.tabs(["📝 核心請求項翻譯", "🎯 11大天條戰略解析"])
            with ta:
                with st.container(border=True, height=650):
                    st.success(st.session_state.m6_report_data.get('translation', '無翻譯結果。'))
            with tb:
                with st.container(border=True, height=650):
                    st.markdown(st.session_state.m6_report_data.get('ip_report', '無解析報告。'))
