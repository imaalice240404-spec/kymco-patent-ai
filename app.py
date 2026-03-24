import os
import json
import random
import time
import datetime
import tempfile
import io
import base64
import streamlit as st
import pandas as pd
import numpy as np
import re
import google.generativeai as genai
from supabase import create_client, Client
import pypdfium2 as pdfium
from docx import Document
import streamlit.components.v1 as components
import plotly.express as px
from PIL import Image, ImageOps

# ==========================================
# ⚙️ 1. 系統初始化與金鑰設定
# ==========================================
st.set_page_config(page_title="機車專利 AI 戰略分析系統", layout="wide")

api_keys = [
    st.secrets.get("GOOGLE_API_KEY_1", st.secrets.get("GOOGLE_API_KEY", "")),
    st.secrets.get("GOOGLE_API_KEY_2", st.secrets.get("GOOGLE_API_KEY", ""))
]
selected_key = random.choice([k for k in api_keys if k])
if selected_key:
    genai.configure(api_key=selected_key)

model = genai.GenerativeModel(
    'gemini-2.5-flash',
    generation_config=genai.types.GenerationConfig(temperature=0.1, top_p=0.8)
)

@st.cache_resource
def init_supabase() -> Client:
    url = st.secrets["SUPABASE_URL"]
    key = st.secrets["SUPABASE_KEY"]
    return create_client(url, key)

supabase = init_supabase()

# --- 狀態追蹤 ---
if 'selected_patents_set' not in st.session_state: st.session_state.selected_patents_set = set()
if 'target_single_patent' not in st.session_state: st.session_state.target_single_patent = None
if 'ip_report_content' not in st.session_state: st.session_state.ip_report_content = ""
if 'rd_card_data' not in st.session_state: st.session_state.rd_card_data = None
if 'claim_data_t2' not in st.session_state: st.session_state.claim_data_t2 = None
if 'pdf_bytes_main' not in st.session_state: st.session_state.pdf_bytes_main = None
if 'scanned_pages' not in st.session_state: st.session_state.scanned_pages = {}
if 'thumbnail_base64' not in st.session_state: st.session_state.thumbnail_base64 = None

# --- 宏觀分析暫存 ---
if 'target_macro_pool' not in st.session_state: st.session_state.target_macro_pool = pd.DataFrame()
if 'ai_macro_matrix' not in st.session_state: st.session_state.ai_macro_matrix = None

# --- 模組五分析暫存 ---
if 'm5_result' not in st.session_state: st.session_state.m5_result = None

# ==========================================
# 🛠️ 2. 輔助函數
# ==========================================
def safe_str(val):
    if pd.isna(val) or val is None: return ""
    return str(val).strip()

def clean_assignee(name):
    name = safe_str(name)
    if not name: return "未知"
    name = re.split(r'股份有限公司|有限公司|公司', name)[0].strip()
    name = name.split(' ')[0].strip() 
    return name if name else "未知"

DB_COL_MAP = {
    'id': 'ID', 'app_num': '申請號', 'cert_num': '證書號', 'pub_date': '公開公告日',
    'assignee': '專利權人', 'title': '專利名稱', 'abstract': '摘要', 'claims': '請求項',
    'legal_status': '案件狀態', 'status': '狀態', 'sys_main': '五大類', 'sys_sub': '次系統',
    'mechanism': '特殊機構', 'effect': '達成功效', 'solution': '核心解法',
    'thumbnail_base64': '代表圖', 'ipc': 'IPC' 
}

def fetch_patents_from_db(status_filter=None):
    query = supabase.table('patents').select("*")
    if status_filter: query = query.eq('status', status_filter)
    response = query.execute()
    df = pd.DataFrame(response.data)
    if not df.empty: df = df.rename(columns=DB_COL_MAP)
    return df

def crop_white_margins(img):
    if img.mode != 'RGB': img = img.convert('RGB')
    inv = ImageOps.invert(img)
    bbox = inv.getbbox()
    if bbox:
        padded_bbox = (max(0, bbox[0]-20), max(0, bbox[1]-20), min(img.width, bbox[2]+20), min(img.height, bbox[3]+20))
        return img.crop(padded_bbox)
    return img

def generate_thumbnail_base64(pdf_bytes, page_num=2, max_size=800):
    try:
        idx = page_num - 1
        doc = pdfium.PdfDocument(pdf_bytes)
        if idx < 0 or idx >= len(doc): idx = 0 
        
        page = doc[idx]
        pil_image = page.render(scale=2.0).to_pil()
        pil_image = crop_white_margins(pil_image)
        
        if hasattr(Image, 'Resampling'):
            pil_image.thumbnail((max_size, max_size), Image.Resampling.LANCZOS)
        else:
            pil_image.thumbnail((max_size, max_size), Image.ANTIALIAS)
            
        buffered = io.BytesIO()
        pil_image.save(buffered, format="JPEG", quality=85)
        img_str = base64.b64encode(buffered.getvalue()).decode()
        return img_str
    except Exception as e:
        return None

def create_word_doc(text):
    doc = Document()
    doc.add_heading('專利戰略深度分析報告', 0)
    for para in text.split('\n'):
        if para.strip(): doc.add_paragraph(para.strip())
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ==========================================
# 📊 3. 側邊欄與標題
# ==========================================
with st.sidebar:
    st.markdown("### ☁️ 雲端資料庫狀態")
    res_all = supabase.table('patents').select("id", count="exact").execute()
    st.info(f"🗄️ 雲端總筆數: {res_all.count if res_all.count else 0}")
    st.markdown("---")
    if st.button("🗑️ 重置單篇暫存 (清除當前分析)", use_container_width=True):
        st.session_state.target_single_patent = None
        st.session_state.pdf_bytes_main = None
        st.session_state.ip_report_content = ""
        st.session_state.rd_card_data = None
        st.session_state.claim_data_t2 = None
        st.session_state.scanned_pages = {}
        st.session_state.thumbnail_base64 = None
        st.rerun()

    st.markdown("---")
    st.markdown("⚠️ **開發者選項：**")
    if st.button("🚨 清空雲端資料庫 (危險)", use_container_width=True):
        supabase.table('patents').delete().neq('status', 'NONE').execute()
        st.success("已清空！請重新匯入。")
        time.sleep(1)
        st.rerun()

st.title("🏍️ 機車專利 AI 戰略分析系統")

tab_ingest, tab_dashboard, tab_single, tab_macro, tab_combine = st.tabs([
    "📥 模組一：雲端探勘匯入", 
    "📊 模組二：研發知識庫", 
    "🕵️ 模組三：單篇深度拆解", 
    "🗺️ 模組四：傳統宏觀地圖",
    "⚔️ 模組五：組合核駁分析"
])

# ==========================================
# 模組一：雲端探勘與資料匯入
# ==========================================
with tab_ingest:
    st.header("1. TWPAT 資料匯入與狀態更新")
    uploaded_excel = st.file_uploader("上傳 TWPAT 匯出的 Excel/CSV (請另存為 .xlsx 格式)", type=["xlsx", "xls", "csv"])

    if uploaded_excel:
        if st.button("🔄 執行資料比對與匯入/更新", type="primary"):
            df = pd.read_csv(uploaded_excel) if uploaded_excel.name.endswith('.csv') else pd.read_excel(uploaded_excel)
            col_map = {
                'title': next((c for c in df.columns if '名稱' in c or '標題' in c), None),
                'abs': next((c for c in df.columns if '摘要' in c), None),
                'claim': next((c for c in df.columns if '範圍' in c or '請求' in c), None),
                'app_num': next((c for c in df.columns if '申請號' in c), None),
                'cert_num': next((c for c in df.columns if '證書' in c or '公告號' in c or '公開號' in c), None),
                'date': next((c for c in df.columns if '日' in c and ('公開' in c or '公告' in c)), None),
                'assignee': next((c for c in df.columns if '權人' in c or '申請人' in c), None),
                'status': next((c for c in df.columns if '狀態' in c), None),
                'ipc': next((c for c in df.columns if 'IPC' in c.upper()), None)
            }
            
            existing_data = supabase.table('patents').select("id, app_num, legal_status").execute()
            existing_dict = {d['app_num']: d for d in existing_data.data if d['app_num']}

            new_rows_to_insert = []
            update_count = 0
            skip_records = 0
            progress_bar_import = st.progress(0)
            
            for i, row in df.iterrows():
                app_val = safe_str(row[col_map['app_num']]) if col_map['app_num'] else ""
                cert_val = safe_str(row[col_map['cert_num']]) if col_map['cert_num'] else ""
                new_status = safe_str(row[col_map['status']]) if col_map['status'] else "未知"
                
                if not app_val: continue 
                
                if app_val in existing_dict:
                    old_record = existing_dict[app_val]
                    old_status = old_record['legal_status']
                    
                    if old_status != new_status:
                        update_payload = {'legal_status': new_status}
                        if "公告" in new_status or "核准" in new_status:
                            update_payload['cert_num'] = cert_val
                            update_payload['claims'] = safe_str(row[col_map['claim']]).replace('\n', '')[:500] if col_map['claim'] else "無請求項"
                            update_payload['rd_card_json'] = None
                            update_payload['vis_data_json'] = None
                            update_payload['ip_report_text'] = None
                            update_payload['thumbnail_base64'] = None
                        
                        supabase.table('patents').update(update_payload).eq('id', old_record['id']).execute()
                        update_count += 1
                    else:
                        skip_records += 1 
                else:
                    new_row = {
                        'app_num': app_val, 'cert_num': cert_val,
                        'pub_date': safe_str(row[col_map['date']]) if col_map['date'] else "未知",
                        'assignee': clean_assignee(safe_str(row[col_map['assignee']])),
                        'title': safe_str(row[col_map['title']]) if col_map['title'] else "無名稱",
                        'abstract': safe_str(row[col_map['abs']]).replace('\n', '')[:500] if col_map['abs'] else "無摘要",
                        'claims': safe_str(row[col_map['claim']]).replace('\n', '')[:500] if col_map['claim'] else "無請求項",
                        'legal_status': new_status,
                        'status': 'PENDING',
                        'ipc': safe_str(row[col_map['ipc']]) if col_map['ipc'] else "未知" 
                    }
                    new_rows_to_insert.append(new_row)
                    
                if i % 10 == 0: progress_bar_import.progress(min(1.0, (i + 1) / len(df)))
            
            progress_bar_import.progress(1.0)
            
            if len(new_rows_to_insert) > 0:
                for i in range(0, len(new_rows_to_insert), 500):
                    supabase.table('patents').insert(new_rows_to_insert[i:i+500]).execute()
            st.success(f"✅ 雲端同步完成！\n- 🆕 新增：{len(new_rows_to_insert)} 筆\n- 🔄 狀態更新：{update_count} 筆\n- ⏭️ 無變動跳過：{skip_records} 筆")

    st.markdown("---")
    st.header("2. AI 批次特徵萃取 (直接更新雲端)")
    
    failed_df = fetch_patents_from_db('FAILED')
    if not failed_df.empty:
        st.warning(f"⚠️ 系統偵測到有 {len(failed_df)} 筆專利先前 AI 解析失敗。")
        if st.button("🔄 點此將失敗專利重新加入排隊 (重置為待處理)"):
            with st.spinner("正在重置狀態..."):
                for _, row in failed_df.iterrows():
                    supabase.table('patents').update({'status': 'PENDING'}).eq('id', row['ID']).execute()
            st.success("✅ 已全數重新加入排隊！畫面即將重整...")
            time.sleep(1)
            st.rerun()

    pending_df = fetch_patents_from_db('PENDING')
    if not pending_df.empty:
        max_pending = len(pending_df)
        if max_pending > 1:
            batch_size = st.slider("選擇處理筆數", 1, min(50, max_pending), min(5, max_pending))
        else:
            batch_size = 1
            st.info("💡 目前僅有 1 筆新資料待處理")
            
        if st.button(f"🤖 啟動高階探勘管線 (處理 {batch_size} 筆)", type="primary"):
            process_df = pending_df.head(batch_size)
            progress_bar = st.progress(0)
            status_text = st.empty()
            for i, (idx, row) in enumerate(process_df.iterrows()):
                status_text.text(f"正在分析 ({i+1}/{batch_size}): {row['專利名稱']} ...")
                
                # 淨化後的字串，無表情符號，使用雙引號
                prompt = f"""
                你是一位具備 20 年經驗的機車廠資深研發主管(RD)兼專利工程師。
                【請嚴格輸出 JSON 格式】：
                {{
                  "五大類": "【最高嚴格限制】絕對只能從這 6 個詞彙中挑選：[動力引擎, 車架懸吊, 電裝, 機電, 車體外觀, 其他]。禁止發明新詞，禁止使用斜線。若你覺得是『懸吊』或『車架』，必須強制寫成『車架懸吊』。若都不符合，請強制寫『其他』。可多選，用半形逗號分隔。",
                  "次系統": "自訂 5-8 字的具體系統名",
                  "特殊機構": "15字內精準描述其物理改變",
                  "達成功效": "20字內描述解決的痛點",
                  "核心解法": "用 RD 聽得懂的白話文，精確描述零件之間的連接與作動關係。"
                }}
                【待分析專利】：
                【名稱】：{row['專利名稱']}
                【摘要】：{row['摘要']}
                【請求項】：{row['請求項']}
                """
                try:
                    res = model.generate_content(prompt)
                    cln = res.text.replace('```json', '').replace('```', '').strip()
                    res_json = json.loads(cln[cln.find('{'):cln.rfind('}')+1])
                    
                    raw_cat = res_json.get('五大類', '其他')
                    valid_cats = ['動力引擎', '車架懸吊', '電裝', '機電', '車體外觀', '其他']
                    final_cats = [c.strip() for c in raw_cat.split(',') if c.strip() in valid_cats]
                    if not final_cats: final_cats = ['其他']
                    
                    supabase.table('patents').update({
                        'sys_main': ', '.join(final_cats), 
                        'sys_sub': res_json.get('次系統', '未分類'),
                        'mechanism': res_json.get('特殊機構', ''), 'effect': res_json.get('達成功效', ''),
                        'solution': res_json.get('核心解法', ''), 'status': 'COMPLETED'
                    }).eq('id', row['ID']).execute()
                except Exception as e:
                    supabase.table('patents').update({'status': 'FAILED'}).eq('id', row['ID']).execute()
                progress_bar.progress((i + 1) / batch_size)
                time.sleep(4)
            st.success("✅ 批次分析完成！")
            time.sleep(1)
            st.rerun()

# ==========================================
# 模組二：研發知識庫與任務分發
# ==========================================
with tab_dashboard:
    completed_df = fetch_patents_from_db('COMPLETED')
    if completed_df.empty:
        st.warning("⚠️ 目前無已分析的資料，請先至模組一匯入。")
    else:
        completed_df['專利類型'] = completed_df.apply(
            lambda x: '發明專利 (I)' if str(x['證書號'] if x['證書號'] else x['申請號']).strip().upper().startswith('I') else
                      ('新型專利 (M)' if str(x['證書號'] if x['證書號'] else x['申請號']).strip().upper().startswith('M') else
                      ('設計專利 (D)' if str(x['證書號'] if x['證書號'] else x['申請號']).strip().upper().startswith('D') else '其他')),
            axis=1
        )

        st.header("🔍 研發技術情報檢索 (R&D Filter Hub)")
        with st.container(border=True):
            col_f1, col_f2, col_f3 = st.columns(3)
            all_main_cats = pd.Series([cat.strip() for cats in completed_df['五大類'].astype(str) for cat in cats.split(',') if cat.strip()]).unique()
            filter_main = col_f1.multiselect("📂 1. 技術系統分類 (支援跨部門搜尋)", list(all_main_cats))
            filter_sub = col_f2.text_input("⚙️ 2. 次系統 (精確搜尋)")
            search_query = col_f3.text_input("🔑 3. 關鍵字 (痛點/解法/機構)")

            col_f4, col_f5, col_f6 = st.columns(3)
            filter_company = col_f4.multiselect("🏢 4. 競爭對手", [c for c in completed_df['專利權人'].unique() if c and c != "未知"])
            filter_type = col_f5.multiselect("📜 5. 專利類型", ["發明專利 (I)", "新型專利 (M)", "設計專利 (D)"])
            filter_status = col_f6.multiselect("⚖️ 6. 法律狀態", ["🔴 有效專利 (公告/核准)", "🟡 審查中 (公開)", "🟢 開源/失效 (消滅/撤回/無效)"])
            
            temp_dates = pd.to_datetime(completed_df['公開公告日'], errors='coerce')
            min_val = temp_dates.min().date() if not pd.isna(temp_dates.min()) else datetime.date(2000, 1, 1)
            filter_date = st.date_input("📅 7. 公開/公告日區間", value=(min_val, datetime.date.today()))

        filtered_df = completed_df.copy()
        
        if filter_main:
            cat_mask = pd.Series([False] * len(filtered_df), index=filtered_df.index)
            for cat in filter_main:
                cat_mask |= filtered_df['五大類'].astype(str).str.contains(cat, na=False)
            filtered_df = filtered_df[cat_mask]
            
        if filter_sub: filtered_df = filtered_df[filtered_df['次系統'].astype(str).str.contains(filter_sub, na=False)]
        if filter_company: filtered_df = filtered_df[filtered_df['專利權人'].apply(lambda x: any(c in str(x) for c in filter_company))]
        if filter_type: filtered_df = filtered_df[filtered_df['專利類型'].isin(filter_type)]
        
        if filter_status:
            status_mask = pd.Series([False] * len(filtered_df), index=filtered_df.index)
            for s in filter_status:
                if "有效" in s: status_mask |= filtered_df['案件狀態'].astype(str).str.contains("公告|核准")
                elif "審查中" in s: status_mask |= filtered_df['案件狀態'].astype(str).str.contains("公開") & ~filtered_df['案件狀態'].astype(str).str.contains("公告|核准")
                elif "開源/失效" in s: status_mask |= filtered_df['案件狀態'].astype(str).str.contains("消滅|撤回|無效|核駁")
            filtered_df = filtered_df[status_mask]

        if len(filter_date) == 2:
            valid_mask = pd.to_datetime(filtered_df['公開公告日'], errors='coerce').dt.date.between(filter_date[0], filter_date[1])
            filtered_df = filtered_df[valid_mask | (filtered_df['公開公告日'] == "未知")]
        if search_query: filtered_df = filtered_df[filtered_df.astype(str).apply(lambda x: x.str.contains(search_query, case=False)).any(axis=1)]

        st.info(f"✨ 檢索出 **{len(filtered_df)}** 筆專利。")

        selected_ids = [k.split("chk_")[1] for k, v in st.session_state.items() if k.startswith("chk_") and v]
        col_act1, col_act2 = st.columns(2)
        with col_act1:
            if st.button(f"🗺️ 將下方 {len(filtered_df)} 筆專利送往【傳統宏觀地圖】分析", use_container_width=True, type="primary"):
                st.session_state.target_macro_pool = filtered_df
                st.session_state.ai_macro_matrix = None
                st.toast("✅ 已傳送！請點擊上方『模組四』頁籤檢視圖表。")
        with col_act2:
            if st.button(f"⚔️ 切換至【模組五】進行組合分析", use_container_width=True, type="secondary"):
                st.toast("👉 請點擊上方『模組五』頁籤開啟分析引擎！")
        
        st.markdown("---")

        for _, p in filtered_df.iterrows():
            disp_id = p['證書號'] if p['證書號'] else p['申請號']
            
            with st.container(border=True):
                col_chk, col_thumb, col_content = st.columns([0.5, 2.5, 7])
                
                with col_chk:
                    st.write("")
                    st.checkbox("選取", key=f"chk_{disp_id}", label_visibility="collapsed")
                
                with col_thumb:
                    st.write("")
                    thumb_b64 = p.get('代表圖')
                    if thumb_b64 and len(str(thumb_b64)) > 100:
                        st.image(f"data:image/jpeg;base64,{thumb_b64}", use_container_width=True)
                    else:
                        st.markdown("<div style='border:1px dashed #ccc; height:180px; display:flex; align-items:center; justify-content:center; color:#999; border-radius:8px; text-align:center;'>🖼️ 待深度拆解<br>解鎖代表圖</div>", unsafe_allow_html=True)
                
                with col_content:
                    if "消滅" in str(p['案件狀態']) or "撤回" in str(p['案件狀態']) or "無效" in str(p['案件狀態']):
                        st.markdown(f"#### 🟢 [{disp_id}] {p['專利名稱']} (開源技術)")
                    elif "公開" in str(p['案件狀態']) and "公告" not in str(p['案件狀態']):
                        st.markdown(f"#### 🟡 [{disp_id}] {p['專利名稱']} (審查中)")
                    else:
                        st.markdown(f"#### 🔴 [{disp_id}] {p['專利名稱']}")
                    
                    st.caption(f"🏢 權利人: {p['專利權人']} ｜ 📅 日期: {p['公開公告日']} ｜ ⚖️ 狀態: {p['案件狀態']} ｜ 🏷️ 類型: **{p['專利類型']}**")
                    
                    c1, c2, c3 = st.columns([2, 1.5, 2])
                    c1.info(f"📂 **系統分類**：\n{p['五大類']} ➡️ {p['次系統']}")
                    c2.warning(f"⚙️ **特殊機構**：\n{p['特殊機構']}")
                    c3.error(f"🎯 **達成功效**：\n{p['達成功效']}")
                    
                    solution_text = p.get('核心解法', '')
                    if solution_text:
                        st.write("")
                        st.markdown(f"<div style='border-left: 4px solid #ddd; padding-left: 15px; color: #555; line-height: 1.6; font-size: 15px;'>💡 **核心解法：**{solution_text}</div>", unsafe_allow_html=True)
                    
                    st.markdown("<br>", unsafe_allow_html=True)
                    if st.button("📄 進入單篇深度拆解 (解鎖圖面)", key=f"btn_s_{disp_id}", use_container_width=True):
                        st.session_state.target_single_patent = p.to_dict()
                        st.session_state.pdf_bytes_main = None 
                        st.session_state.thumbnail_base64 = None 
                        st.balloons()
                        st.success(f"✅ 已鎖定 [{disp_id}]！請點擊上方 **「🕵️ 模組三：單篇深度拆解」** 頁籤。")

# ==========================================
# 🌟 模組三：單篇深度拆解工作站
# ==========================================
with tab_single:
    if not st.session_state.target_single_patent:
        st.warning("👈 請先至【模組二：研發知識庫】挑選一篇專利，並點擊「進入單篇深度拆解」。")
    else:
        target = st.session_state.target_single_patent
        target_id = target.get('證書號') or target.get('申請號')
        db_id = target.get('ID') 
        is_utility_model = str(target_id).upper().startswith('M')
        applicant_main = target.get('專利權人', '未知')
        
        latest_res = supabase.table('patents').select("rd_card_json, vis_data_json, ip_report_text, thumbnail_base64").eq('id', db_id).execute()
        has_cache = False
        has_thumb = False
        if latest_res.data and latest_res.data[0].get('rd_card_json'):
            has_cache = True
            cached_data = latest_res.data[0]
            if latest_res.data[0].get('thumbnail_base64'):
                has_thumb = True
        
        st.header(f"🕵️ 單篇深度拆解：[{target_id}] {target.get('專利名稱')}")
        
        clean_num_m = ''.join(e for e in str(target_id) if e.isalnum())
        gpt_url = f"https://patents.google.com/patent/TW{clean_num_m}B" if not is_utility_model else f"https://patents.google.com/patent/TW{clean_num_m}U"
        
        st.markdown(f"**🏢 權利人：** {applicant_main} | **📅 公開日：** {target.get('公開公告日')} ｜ 🏷️ 類型：{target.get('專利類型')} 👉 [Google Patents 傳送門]({gpt_url})")
        
        st.markdown("---")
        
        with st.container(border=True):
            st.subheader("📥 第一步：請上傳本案之 PDF 說明書以啟動雙向連動")
            
            uploaded_pdf = st.file_uploader("上傳 PDF", type=["pdf"], key="pdf_deep_dive")
            col_pg, _ = st.columns([2, 2])
            with col_pg: target_fig_page = st.number_input("🖼️ 指定專利代表圖所在頁碼", min_value=1, value=2, key="single_rep_fig_pg")

            btn_text = "⚡ 瞬間解鎖單篇深度拆解 (雲端快取)" if has_cache else "🚀 啟動單篇 AI 深度解剖 (消耗 API)"
            
            if st.button(btn_text, type="primary", use_container_width=True):
                if not uploaded_pdf: st.error("⚠️ 請上傳 PDF！")
                else:
                    st.session_state.pdf_bytes_main = uploaded_pdf.getvalue()
                    st.session_state.scanned_pages = {} 
                    generated_b64 = None
                    if not has_thumb: generated_b64 = generate_thumbnail_base64(st.session_state.pdf_bytes_main, page_num=target_fig_page)
                    
                    if has_cache:
                        st.session_state.rd_card_data = cached_data['rd_card_json']
                        st.session_state.claim_data_t2 = cached_data['vis_data_json']
                        st.session_state.ip_report_content = cached_data['ip_report_text']
                        if generated_b64:
                            st.session_state.thumbnail_base64 = generated_b64
                            supabase.table('patents').update({ 'thumbnail_base64': generated_b64 }).eq('id', db_id).execute()
                    else:
                        with st.spinner("🧠 AI 正在地毯式搜索全文 (約 30-40 秒)..."):
                            try:
                                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
                                    tmp_file.write(st.session_state.pdf_bytes_main)
                                    tmp_file_path = tmp_file.name
                                gemini_file = genai.upload_file(tmp_file_path)
                                
                                ip_report_template = """
【一、 FTO 風險判定】
(紅燈：具威脅 / 黃燈：需注意 / 綠燈：已失效。並簡述判定與證書號)
防呆原則：依「目前狀態」判斷，忽略日期推算。若狀態為「公告/核准」或「公開」，絕對不可判定為綠燈！若為「消滅/無效」或「撤回」才可判為綠燈。

【二、 技術核心快照】
1. 發明目的： (說明解決傳統弊病)
2. 核心技術： (說明具體零件結構設計)
3. 宣稱功效： (說明提升了什麼物理效果)

【三、 研發部門精準派發】
[填入建議部門]。 (分發理由)

【四、 先前技術與妥協分析】
本案欲解決之舊設計缺點： (習用技術缺點)
空間配置限制： (列出獨立項限縮最嚴格之特徵)

【五、 獨立項全要件拆解】
最廣獨立項（請求項1）拆解：
(以 1. 2. 3. 逐行條列拆解，不要加註解！)
破口： (精準點出最容易被迴避的限制條件)

【六、 附屬項隱藏地雷探測】
(條列出具備具體結構形狀、位置、或參數限制的附屬項)

【七、 侵權可偵測性評估】
(極易偵測 / 需破壞性拆解，並給出理由)

【八、 實證功效檢驗】
(是否有實體測試數據，或僅為定性描述)

【九、 高階迴避設計建議】
(提出基於破口的具體修改機構方向)

【十、 技術演進與機構整併雷達】
(分析屬於機構整併或架構重組，並說明解決了什麼困境)
                                """

                                # 淨化字串，無表情符號，使用雙引號
                                prompt_master = f"""
                                【資深機車專利主管語氣】：請仔細閱讀 PDF 檔案。
                                【輸出格式嚴格要求：純 JSON 格式】
                                {{
                                  "rd_card": {{
                                    "title": "一句話總結", 
                                    "problem": "傳統缺點", 
                                    "solution": "本專利特殊結構",
                                    "risk_check": ["1-1. 獨立項全要件特徵A", "1-2. 獨立項全要件特徵B"],
                                    "design_avoid_rd": ["針對獨立項限制A的迴避方向", "針對獨立項限制B的迴避方向"]
                                  }},
                                  "vis_data": {{
                                    "claims": ["1. 獨立項全文...", "2. 依據請求項1..."],
                                    "components": [ {{"id": "10", "name": "車架"}} ],
                                    "spec_texts": ["段落內容全文"],
                                    "loophole_quote": "請直接從上方請求項原文中，一字不漏複製最能代表本案特徵或破口的那一段文字。請勿包含習知技術，且連標點符號都必須與原文一致，否則系統無法上色！"
                                  }},
                                  "ip_report": "請填寫完下方 IP報告十點 後輸出在此，不要使用 Markdown 格式。"
                                }}
                                【補充指示】：
                                1. rd_card.risk_check 請務必「逐項拆解請求項 1 (獨立項) 的所有全要件限制」。
                                2. vis_data.claims 請務必保留「請求項全文的數字編號」，絕對不可省略。
                                3. vis_data.components 請【極度精確】萃取請求項出現的元件與標號。絕對不可張冠李戴配錯對（例如說明書寫第一管部22相當於下降管部，則下降管部的 id 就是 22，絕不可誤植為 23）。

                                【IP報告結構】：
                                {ip_report_template}
                                """
                                response = model.generate_content([gemini_file, prompt_master])
                                clean_text = response.text.replace('```json', '').replace('```', '').strip()
                                clean_text = clean_text[clean_text.find('{'):clean_text.rfind('}')+1]
                                master_json = json.loads(clean_text)

                                st.session_state.rd_card_data = master_json.get("rd_card")
                                st.session_state.claim_data_t2 = master_json.get("vis_data")
                                st.session_state.ip_report_content = master_json.get("ip_report")
                                
                                supabase.table('patents').update({
                                    'rd_card_json': st.session_state.rd_card_data,
                                    'vis_data_json': st.session_state.claim_data_t2,
                                    'ip_report_text': st.session_state.ip_report_content,
                                    'thumbnail_base64': generated_b64
                                }).eq('id', db_id).execute()
                                os.remove(tmp_file_path)
                                genai.delete_file(gemini_file.name)
                            except Exception as e: st.error(f"分析失敗：{e}")

        st.markdown("<br>", unsafe_allow_html=True)

        if st.session_state.rd_card_data and st.session_state.pdf_bytes_main:
            sub_tab_rd, sub_tab_ip = st.tabs(["🧑‍💻 Tab 1 研發：迴避設計大屏", "⚖️ Tab 2 智權：法務審查中心"])
            
            with sub_tab_rd:
                rd_data = st.session_state.rd_card_data
                
                col_c1, col_c2, col_c3 = st.columns([1.5, 2, 1.5])
                with col_c1:
                    with st.container(border=True, height=480):
                        st.markdown(f"#### 🎯 研發戰略看板\n**{rd_data.get('title', '未知')}**")
                        st.markdown(f"**🔥 解決痛點：** {rd_data.get('problem', '')}\n**💡 核心解法：** {rd_data.get('solution', '')}")

                with col_c2:
                    with st.container(border=True, height=480):
                        st.markdown("#### 🛡️ 獨立項（最廣範圍）全要件檢核")
                        st.caption("全要件原則：若我司設計符合下方【所有】特徵，則侵權風險極高。")
                        risk_list = rd_data.get('risk_check', [])
                        checked_count = 0
                        for i, risk in enumerate(risk_list):
                            if st.checkbox(f"{risk}", key=f"risk_c_{i}"): checked_count += 1
                        
                        st.markdown("<br>", unsafe_allow_html=True)
                        if len(risk_list) > 0:
                            if checked_count == len(risk_list): st.markdown("<div style='padding:10px; background-color:#ffebee; color:#c62828; border-radius:5px;'><b>⚠️ 警告：全要件命中，高度侵權風險！</b></div>", unsafe_allow_html=True)
                            else: st.markdown("<div style='padding:10px; background-color:#e8f5e9; color:#2e7d32; border-radius:5px;'><b>🎉 至少一要件未命中，文義迴避成功。</b></div>", unsafe_allow_html=True)

                with col_c3:
                    with st.container(border=True, height=480):
                        st.markdown("#### 🛡️ 高階迴避建議方向")
                        st.caption("建議研發從下方機構特徵進行「實質修改」：")
                        for avoid in rd_data.get('design_avoid_rd', []): st.markdown(f"✅ {avoid}")

                st.markdown("---")
                st.markdown("### 🎯 終極雙向連動大屏 (支援圖面旋轉)")
                
                pdf_doc_v = pdfium.PdfDocument(st.session_state.pdf_bytes_main)
                total_pages_v = len(pdf_doc_v)

                col_page, col_rot, col_btn = st.columns([1, 1.5, 1.5])
                with col_page:
                    target_page = st.number_input(f"📄 圖紙頁碼 (共 {total_pages_v} 頁)", min_value=1, max_value=total_pages_v, value=min(2, total_pages_v), key="vis_page_rd")
                with col_rot:
                    rot_angle = st.radio("🔄 圖面旋轉", [0, 90, 180, 270], horizontal=True, key="vis_rot_angle")
                
                page = pdf_doc_v[target_page - 1]
                raw_pil_img = page.render(scale=2.0).to_pil()
                
                if rot_angle != 0:
                    raw_pil_img = raw_pil_img.rotate(-rot_angle, expand=True, fillcolor='white')
                
                cropped_img = crop_white_margins(raw_pil_img) 
                img_byte_arr = io.BytesIO()
                cropped_img.save(img_byte_arr, format='JPEG')
                encoded_img = base64.b64encode(img_byte_arr.getvalue()).decode()
                img_uri = f"data:image/jpeg;base64,{encoded_img}"

                scan_key = f"{target_page}_{rot_angle}"
                is_scanned = scan_key in st.session_state.scanned_pages
                
                with col_btn:
                    st.write("")
                    if not is_scanned:
                        if st.button(f"🔍 啟動圖片標號鎖定", use_container_width=True, key="btn_scan_rd"):
                            with st.spinner("Gemini Vision 正在極高精度鎖定座標..."):
                                try:
                                    comp_dict_list = st.session_state.claim_data_t2.get("components", [])
                                    known_comps_str = json.dumps(comp_dict_list, ensure_ascii=False)
                                    
                                    # 淨化字串
                                    prompt_vision = f"""這是一張專利圖。已知元件表：{known_comps_str}。
                                    請找出圖片上「所有肉眼可見的數字標號」，並精準估算其「幾何中心點」的相對座標(x_rel, y_rel，範圍0.000~1.000，請精確到小數點後三位)。
                                    【極度要求】：座標必須極度精準地對準數字的正中心！如果該頁「無標號」或是「純文字」，請輸出空的陣列：{{ "hotspots": [] }}。
                                    嚴格輸出 JSON 格式。範例：{{ "hotspots": [ {{"number": "31", "name": "汽缸頭", "x_rel": 0.452, "y_rel": 0.551}} ] }}"""
                                    
                                    response_vis = model.generate_content([cropped_img, prompt_vision])
                                    if not response_vis.text: ai_visual_data = []
                                    else:
                                        clean_text_vis = response_vis.text.replace('```json', '').replace('```', '').strip()
                                        start_idx = clean_text_vis.find('{')
                                        end_idx = clean_text_vis.rfind('}')
                                        if start_idx != -1 and end_idx != -1:
                                            clean_text_vis = clean_text_vis[start_idx:end_idx+1]
                                            ai_visual_data = json.loads(clean_text_vis).get("hotspots", [])
                                        else: ai_visual_data = [] 
                                    st.session_state.scanned_pages[scan_key] = ai_visual_data
                                    st.rerun()
                                except Exception as e: st.error(f"視覺解析失敗：{e}")
                    else:
                        if not st.session_state.scanned_pages[scan_key]: st.warning("⚡ 掃描完成，未偵測到標號。")
                        else: st.success("⚡ 座標已鎖定！體驗下方雙向連動。")

                claim_lines = st.session_state.claim_data_t2.get("claims", [])
                comp_dict_list = st.session_state.claim_data_t2.get("components", [])
                loophole_quote = st.session_state.claim_data_t2.get("loophole_quote", "")
                
                final_claims_html_list = []
                for i, line in enumerate(claim_lines):
                    processed_line = line
                    
                    if i == 0 and loophole_quote and loophole_quote in processed_line:
                        processed_line = processed_line.replace(loophole_quote, f'<mark class="loophole-highlight">{loophole_quote}</mark>')
                    
                    for comp in comp_dict_list:
                        c_num = comp.get("id", "")
                        c_name = comp.get("name", "")
                        replacement = f'<span class="comp-text comp-{c_num}" onmouseover="hoverText(\'{c_num}\')" onmouseout="leaveText(\'{c_num}\')">{c_name} ({c_num})</span>'
                        processed_line = processed_line.replace(f"{c_name} ({c_num})", replacement).replace(c_name, replacement)
                    
                    if i == 0:
                        final_claims_html_list.append(f'<div class="independent-claim-box">{processed_line}</div>')
                    else:
                        final_claims_html_list.append(f'<p class="dependent-claim">{processed_line}</p>')
                
                claim_text_full = "".join(final_claims_html_list)

                hotspots_html = ""
                if is_scanned:
                    for spot in st.session_state.scanned_pages.get(scan_key, []):
                        if spot['name'] != "未知":
                            hotspots_html += f"""
                            <div class="hotspot hotspot-marker-{spot['number']}" id="hotspot-{spot['number']}" style="left: {spot['x_rel']*100}%; top: {spot['y_rel']*100}%;" onmouseover="hoverImage('{spot['number']}', '{spot['name']}')" onmouseout="leaveImage('{spot['number']}')"></div>"""

                css_style = """
                <style>
                    body { margin: 0; font-family: sans-serif; background: #fff; }
                    .main-container { display: flex; height: 800px; width: 100%; border: 1px solid #ddd; border-radius: 8px; overflow: hidden; }
                    .img-section { flex: 6; position: relative; overflow: auto; background: #f8f9fa; border-right: 2px solid #ddd; padding: 10px; display: flex; justify-content: center; align-items: flex-start;}
                    .img-wrapper { position: relative; display: inline-block; }
                    .patent-img { max-width: 100%; height: auto; display: block; }
                    
                    .hotspot { position: absolute; width: 40px; height: 40px; transform: translate(-50%, -50%); border-radius: 50%; cursor: pointer; transition: 0.2s; border: 2px solid transparent; z-index: 10; }
                    .hotspot:hover { background: rgba(255, 0, 0, 0.3); border: 2px solid red; box-shadow: 0 0 10px rgba(255,0,0,0.5); z-index: 50; }
                    .hotspot-active { background: rgba(255, 255, 0, 0.6) !important; border: 3px solid red !important; box-shadow: 0 0 20px red !important; transform: translate(-50%, -50%) scale(1.3); z-index: 50; }
                    #tooltip { display: none; position: absolute; background: rgba(0, 0, 0, 0.8); color: white; padding: 6px 12px; border-radius: 4px; font-size: 14px; z-index: 100; pointer-events: none; white-space: nowrap; }

                    .text-section { flex: 4; padding: 20px; overflow-y: auto; font-size: 16px; line-height: 1.8; color: #333; }
                    
                    .independent-claim-box { background-color: #fafafa; padding: 15px; border-radius: 8px; border-left: 6px solid #94a3b8; margin-bottom: 15px; }
                    
                    .loophole-highlight { background-color: #ffeb3b; font-weight: bold; color: #b45309; padding: 2px 4px; border-radius: 3px; box-shadow: 0 0 5px rgba(255, 235, 59, 0.8); }
                    
                    .dependent-claim { margin-bottom: 15px; color: #555; }
                    .comp-text { color: #0284c7; font-weight: bold; cursor: pointer; border-bottom: 1px dashed #0284c7; padding: 0 2px; transition: 0.2s; }
                    .highlight-active { background-color: #fef08a; color: #b91c1c; border-bottom: none; border-radius: 3px; padding: 2px 4px; }
                </style>
                """
                html_skeleton = f"""
                <!DOCTYPE html><html><head>{css_style}</head><body>
                <div class="main-container">
                    <div class="img-section" id="img-container"><div class="img-wrapper"><img src="{img_uri}" class="patent-img">{hotspots_html}</div><div id="tooltip"></div></div>
                    <div class="text-section"><div style="font-size:18px; font-weight:bold; color:#1e3a8a; margin-bottom:15px; position:sticky; top:0; background:white; z-index:10;">📜 請求項對應 (特徵破口重點標記)</div>{claim_text_full}</div>
                </div>
                <script>
                    const tooltip = document.getElementById('tooltip');
                    function hoverImage(num, name) {{ document.onmousemove = e => {{ tooltip.style.left = (e.pageX + 15) + 'px'; tooltip.style.top = (e.pageY + 15) + 'px'; }}; tooltip.innerHTML = "標號 <b>" + num + "</b> : " + name; tooltip.style.display = 'block'; document.querySelectorAll('.comp-' + num).forEach((el, i) => {{ el.classList.add('highlight-active'); if(i===0) el.scrollIntoView({{behavior:'smooth', block:'center'}}); }}); }}
                    function leaveImage(num) {{ document.onmousemove = null; tooltip.style.display = 'none'; document.querySelectorAll('.comp-' + num).forEach(el => el.classList.remove('highlight-active')); }}
                    function hoverText(num) {{ document.querySelectorAll('.comp-' + num).forEach(el => el.classList.add('highlight-active')); const hs = document.getElementById('hotspot-' + num); if(hs) {{ hs.classList.add('hotspot-active'); hs.scrollIntoView({{behavior:'smooth', block:'center'}}); }} }}
                    function leaveText(num) {{ document.querySelectorAll('.comp-' + num).forEach(el => el.classList.remove('highlight-active')); const hs = document.getElementById('hotspot-' + num); if(hs) hs.classList.remove('hotspot-active'); }}
                </script></body></html>
                """
                components.html(html_skeleton, height=820, scrolling=False)

            with sub_tab_ip:
                st.markdown("## 🏛️ 智權法務審查工作站")
                ip_tab_report, ip_tab_claim = st.tabs(["📄 智權戰略深度報告 (10大天條)", "⚖️ 請求項文義比對 (三視窗)"])
                
                with ip_tab_report:
                    col_r1, col_r2 = st.columns([3, 1])
                    with col_r1: st.markdown("以下為嚴格遵守「智權審查 10 大天條」生成的實務報告：")
                    with col_r2:
                        st.download_button("📥 下載 Word 報告", data=create_word_doc(st.session_state.ip_report_content), file_name=f"IP_Report_{target_id}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                    with st.container(height=650, border=True):
                        st.markdown(st.session_state.ip_report_content)
                
                with ip_tab_claim:
                    components_list = st.session_state.claim_data_t2.get("components", [])
                    if components_list:
                        comp_options = {f"[{c.get('id','')}] {c.get('name','')}": c for c in components_list}
                        col_sel, _ = st.columns([1, 1])
                        with col_sel:
                            selected_comp = st.selectbox(f"🎯 選擇比對目標元件 (共 {len(components_list)} 個)：", list(comp_options.keys()), key="ip_comp_sel")
                            active_c = comp_options[selected_comp]
                        
                        st.markdown("<br>", unsafe_allow_html=True)
                        col_left, col_right = st.columns([1, 1])
                        
                        with col_left:
                            st.markdown("### 🧩 獨立項文義")
                            with st.container(height=350, border=True):
                                for line in st.session_state.claim_data_t2.get("claims", []):
                                    if active_c['name'] in line:
                                        hl_line = line.replace(active_c['name'], f"<span style='background-color:#fff3cd; font-weight:bold; color:#856404; padding:2px 4px; border-radius:3px;'>{active_c['name']}</span>")
                                        st.markdown(f"<div style='padding: 8px; border-bottom: 1px dashed #eee;'>{hl_line}</div>", unsafe_allow_html=True)
                                    else:
                                        st.markdown(f"<div style='padding: 8px; border-bottom: 1px dashed #eee; color: #555;'>{line}</div>", unsafe_allow_html=True)
                            
                            st.markdown("### 🖼️ 專利圖面")
                            pdf_doc_ip = pdfium.PdfDocument(st.session_state.pdf_bytes_main)
                            pg_ip = st.number_input("頁碼", min_value=1, max_value=len(pdf_doc_ip), value=min(2, len(pdf_doc_ip)), key="ip_pg")
                            with st.container(height=450, border=True): 
                                st.image(pdf_doc_ip[pg_ip - 1].render(scale=2.0).to_pil(), use_container_width=True)

                        with col_right:
                            st.markdown("### 📖 說明書具體限制")
                            with st.container(height=895, border=True):
                                st.info(f"📍 目標：**{active_c['name']} ({active_c.get('id','')})**")
                                found_texts = [t for t in st.session_state.claim_data_t2.get('spec_texts', []) if active_c['name'] in t]
                                if not found_texts: st.warning("未找到說明。")
                                else:
                                    for t in found_texts:
                                        hl_t = t.replace(active_c['name'], f"<mark style='background-color:#cce5ff; color:#004085; font-weight:bold; padding:2px; border-radius:3px;'>{active_c['name']}</mark>")
                                        st.markdown(f"<div style='background: #f8f9fa; padding: 10px; border-left: 4px solid #007bff; margin-bottom: 10px;'>{hl_t}</div>", unsafe_allow_html=True)

# ==========================================
# 🗺️ 模組四：傳統宏觀地圖
# ==========================================
with tab_macro:
    st.header("🗺️ 傳統專利大數據分析 (Macro Landscape)")
    
    if st.session_state.target_macro_pool.empty:
        st.warning("👈 請先至【模組二：研發知識庫】設定篩選條件，並點擊「傳統專利分析」按鈕將資料送來這裡。")
    else:
        df_macro = st.session_state.target_macro_pool
        st.success(f"✅ 已成功載入 **{len(df_macro)}** 筆專利進行宏觀分析！")
        
        sub_t1, sub_t2, sub_t3, sub_t4, sub_t6 = st.tabs([
            "🏢 競爭者佈局", "📈 演進趨勢", "🎯 IPC 熱區分析 (四階)", 
            "🧠 AI 技術功效矩陣", "👑 核心地雷探勘"
        ])

        with sub_t1:
            st.markdown("### 🏢 專利權人火力佈局")
            top_app = df_macro['專利權人'].value_counts().reset_index().head(10)
            top_app.columns = ['公司名稱', '專利數量']
            fig_bar = px.bar(top_app, x='專利數量', y='公司名稱', orientation='h', color='專利數量', color_continuous_scale='Blues')
            fig_bar.update_layout(yaxis={'categoryorder':'total ascending'})
            st.plotly_chart(fig_bar, use_container_width=True)

        with sub_t2:
            st.markdown("### 📈 歷年申請/公告趨勢")
            df_macro['年份'] = df_macro['公開公告日'].astype(str).str[:4]
            yt = df_macro['年份'].value_counts().reset_index().sort_values('年份')
            yt.columns = ['年份', '專利數量']
            fig_line = px.line(yt[yt['年份'].str.isnumeric()], x='年份', y='專利數量', markers=True, line_shape='spline', color_discrete_sequence=['#ff7f0e'])
            st.plotly_chart(fig_line, use_container_width=True)

        with sub_t3:
            st.markdown("### 🎯 IPC 熱區分析 (四階)")
            if 'IPC' in df_macro.columns and not df_macro['IPC'].astype(str).replace('未知', '').str.strip().empty:
                df_macro['IPC_四階'] = df_macro['IPC'].apply(lambda x: str(x).split(';')[0].split('|')[0].split('(')[0].strip() if pd.notna(x) and str(x) != "未知" else "未知")
                ipc_d = df_macro[df_macro['IPC_四階'] != '未知']['IPC_四階'].value_counts().reset_index().head(15)
                ipc_d.columns = ['IPC四階', '數量']
                if not ipc_d.empty:
                    fig_pie = px.pie(ipc_d, values='數量', names='IPC四階', hole=0.4)
                    fig_pie.update_traces(textposition='inside', textinfo='percent+label')
                    st.plotly_chart(fig_pie, use_container_width=True)
                else:
                    st.info("沒有足夠的 IPC 資料可供分析。")
            else:
                st.warning("⚠️ 目前的資料庫中沒有 IPC 資料。請依照上述指示更新資料庫後重新匯入 Excel。")

        with sub_t4:
            st.markdown("### 🧠 AI 自動生成：技術功效矩陣")
            analyze_count = st.slider("選擇要投入 AI 矩陣分析的專利數量 (取前 N 筆)", 1, min(len(df_macro), 30), min(len(df_macro), 15), key="slider_macro")
            
            if st.button("🚀 啟動矩陣引擎", use_container_width=True, type="primary"):
                with st.spinner("AI 正在掃描專利機構與功效，繪製戰略地圖..."):
                    try:
                        sample_df = df_macro.head(analyze_count)
                        p_data = "".join([f"[{str(row['證書號'] if row['證書號'] else row['申請號'])}] {str(row['專利名稱'])} | 機構：{str(row['特殊機構'])} | 功效：{str(row['達成功效'])}\n" for _, row in sample_df.iterrows()])
                        
                        # 淨化字串，無表情符號
                        prompt_matrix = f"""
                        請分析以下機車專利資料，並輸出純 JSON 格式。
                        矩陣維度X (達成功效): ["提升散熱與冷卻", "提升燃燒與動力效率", "結構緊湊與輕量化", "降低震動與噪音", "改善潤滑與耐用度", "降低製造成本"]。
                        矩陣維度Y (技術手段): ["汽缸本體與散熱片", "活塞曲軸", "氣門進排氣", "機油道水套", "燃油噴射點火", "引擎外殼", "煞車懸吊", "電控儀表"]。
                        
                        {{
                          "matrix": [{{"專利號": "XXX", "技術手段": "選項", "達成功效": "選項"}}],
                          "top_patents": [{{"專利號": "XXX", "專利名稱": "XXX", "威脅度": "極高/中等", "入選理由": "..."}}]
                        }}
                        資料：{p_data}
                        """
                        res = model.generate_content(prompt_matrix)
                        cln = res.text.replace('```json','').replace('```','').strip()
                        st.session_state.ai_macro_matrix = json.loads(cln[cln.find('{'):cln.rfind('}')+1])
                        st.success("✅ 戰略矩陣解析完成！")
                    except Exception as e: 
                        st.error(f"分析失敗：{e}")

            if st.session_state.ai_macro_matrix:
                fig_heat = px.density_heatmap(pd.DataFrame(st.session_state.ai_macro_matrix["matrix"]), y='技術手段', x='達成功效', text_auto=True, color_continuous_scale='Reds')
                st.plotly_chart(fig_heat, use_container_width=True)

        with sub_t6:
            st.markdown("### 👑 核心地雷探勘 (Killer Patents)")
            if st.session_state.ai_macro_matrix:
                for p in st.session_state.ai_macro_matrix.get("top_patents", []):
                    with st.container(border=True):
                        c = "red" if "高" in p.get("威脅度", "") else "orange"
                        st.markdown(f"#### 🎯 [{p.get('專利號')}] {p.get('專利名稱')}")
                        st.markdown(f"**威脅度：** <span style='color:{c};font-weight:bold;'>{p.get('威脅度')}</span><br>**洞察：** {p.get('入選理由')}", unsafe_allow_html=True)
            else:
                st.info("請先至左側「AI 技術功效矩陣」頁籤啟動分析。")

# ==========================================
# ⚔️ 模組五：組合核駁與進步性分析 (真實攻防版)
# ==========================================
with tab_combine:
    st.header("⚔️ 模組五：組合核駁與進步性分析 (TSM 攻防引擎)")
    st.markdown("本模組基於 TIPO 專利審查基準，結合 **AI 全文理解** 與 **使用者手動輸入之 OA 特徵** 進行雙向推演。")

    completed_df_m5 = fetch_patents_from_db('COMPLETED')
    if completed_df_m5.empty or len(completed_df_m5) < 2:
        st.warning("⚠️ 資料庫中已分析完成的專利不足，請先至模組一匯入引證案。")
    else:
        options = [f"[{row['證書號'] if row['證書號'] else row['申請號']}] {row['專利名稱']}" for _, row in completed_df_m5.iterrows()]
        option_mapping = {f"[{row['證書號'] if row['證書號'] else row['申請號']}] {row['專利名稱']}": row for _, row in completed_df_m5.iterrows()}

        st.markdown("### 📚 第一步：設定背景知識 (自動抓取雲端專利解析)")
        st.caption("AI 將會讀取您選定之引證案的底層技術邏輯，用於尋找『物理衝突』與『破壞發明目的』之防禦彈藥。")
        
        col_m5_bg1, col_m5_bg2 = st.columns(2)
        with col_m5_bg1:
            ref_a_sel = st.selectbox("📄 選擇【引證一】", options, index=0)
        with col_m5_bg2:
            ref_b_sel = st.selectbox("📄 選擇【引證二】", options, index=min(1, len(options)-1))

        st.markdown("### 🎯 第二步：手動標定攻防爭點 (來自 OA 或研發機密)")
        with st.container(border=True):
            target_feature = st.text_area("1. 本案欲保護/答辯之核心特徵 (請貼上被核駁之請求項，或尚未公開之研發特徵)", height=100)
            
            col_m5_text1, col_m5_text2 = st.columns(2)
            with col_m5_text1:
                ref_a_detail = st.text_area("2. 【引證一】官方認定之具體揭露 (請填入 OA 指定之段落/圖式與特徵)", height=150)
            with col_m5_text2:
                ref_b_detail = st.text_area("3. 【引證二】官方認定之具體揭露 (請填入 OA 指定之段落/圖式與特徵)", height=150)
                
            examiner_logic = st.text_area("4. 審查委員之結合邏輯 (委員為何認為兩者具備結合動機？)", height=100)

        if st.button("🚀 啟動 TSM 雙向攻防分析", type="primary", use_container_width=True):
            if not target_feature or not ref_a_detail or not ref_b_detail:
                st.error("⚠️ 請填寫完整的爭點特徵與引證揭露內容，才能進行精準打擊！")
            else:
                with st.spinner("🤖 AI 專利代理人正在調閱背景知識，並針對您輸入的 OA 爭點進行沙盤推演 (約 20-30 秒)..."):
                    ref_a = option_mapping[ref_a_sel]
                    ref_b = option_mapping[ref_b_sel]

                    # 淨化字串，無表情符號
                    prompt_m5 = f"""
                    【角色設定】：你是一位熟悉台灣智財局專利審查基準、具備 20 年經驗的機車領域專利代理人（PHOSITA）。
                    【任務】：請基於使用者提供的「官方 OA 爭點」，並參考兩篇引證專利的全文背景，進行 TSM (結合動機) 雙向攻防推演。

                    【使用者輸入之 OA 爭點與特徵】(你的分析靶心)：
                    - 本案爭點特徵：{target_feature}
                    - 引證一具體揭露 (OA指定)：{ref_a_detail}
                    - 引證二具體揭露 (OA指定)：{ref_b_detail}
                    - 委員結合邏輯：{examiner_logic}

                    【背景知識參考】(用於尋找物理衝突或反向教示)：
                    [引證A背景] 摘要：{ref_a['摘要']} | 核心解法：{ref_a['核心解法']}
                    [引證B背景] 摘要：{ref_b['摘要']} | 核心解法：{ref_b['核心解法']}

                    【請嚴格輸出 JSON 格式進行進步性攻防評估】：
                    {{
                      "delta_feature": "請精準總結引證A缺乏，而由引證B補足的『差異特徵』是什麼？",
                      "attack_argument": {{
                        "conclusion": "結合容易 / 具備進步性核駁空間",
                        "field_problem_match": "分析 A 與 B 在領域與解決問題上的共通性...",
                        "motivation_to_combine": "順著委員的邏輯，論述為何通常知識者有動機將 B 結合至 A..."
                      }},
                      "defense_argument": {{
                        "conclusion": "結合困難 / 違反 Could-Would 測試",
                        "teaching_away": "利用你讀到的背景知識，列出阻礙 A 結合 B 的客觀技術因素（如物理衝突、安裝限制、破壞發明目的等）...",
                        "hindsight_warning": "點出審查委員強行拼湊，犯下了什麼後見之明謬誤..."
                      }}
                    }}
                    """
                    try:
                        res_m5 = model.generate_content(prompt_m5)
                        cln_m5 = res_m5.text.replace('```json', '').replace('```', '').strip()
                        st.session_state.m5_result = json.loads(cln_m5[cln_m5.find('{'):cln_m5.rfind('}')+1])
                        st.success("✅ TSM 雙向攻防推演完成！")
                    except Exception as e:
                        st.error(f"分析失敗，請確認 API 連線或稍後再試。錯誤訊息：{e}")

        if st.session_state.m5_result:
            m5_data = st.session_state.m5_result
            st.markdown("---")
            st.markdown(f"### 🎯 差異特徵 (Delta Feature)\n<div style='background-color:#e0f2fe; color:#0369a1; padding:15px; border-radius:8px; font-weight:bold; font-size:16px; border-left: 5px solid #0284c7;'>{m5_data.get('delta_feature', '')}</div><br>", unsafe_allow_html=True)

            tab_def, tab_atk = st.tabs(["🛡️ 防禦視角 (我方強力答辯)", "⚔️ 攻擊視角 (審查委員邏輯)"])

            with tab_def:
                st.markdown("#### 🛡️ 反向教示與阻礙結合 (Teaching Away)")
                st.error(m5_data.get('defense_argument', {}).get('teaching_away', ''))
                st.markdown("#### ⚠️ 後見之明謬誤 (Hindsight Bias)")
                st.warning(m5_data.get('defense_argument', {}).get('hindsight_warning', ''))

            with tab_atk:
                st.markdown("#### ⚔️ 領域與問題共通性")
                st.info(m5_data.get('attack_argument', {}).get('field_problem_match', ''))
                st.markdown("#### 🔗 結合動機 (Motivation to Combine)")
                st.success(m5_data.get('attack_argument', {}).get('motivation_to_combine', ''))

# ===================== 🚀 程式碼到此完美結束 🚀 =====================
